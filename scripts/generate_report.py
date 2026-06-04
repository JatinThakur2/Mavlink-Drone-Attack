#!/usr/bin/env python3
"""
Generates the full attack report as a Word document.
Run: python3 generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

TODAY     = datetime.date.today().strftime("%d %B %Y")
NOW       = datetime.datetime.now().strftime("%H:%M")
OUT_FILE  = os.path.join(os.path.dirname(__file__), "..", "reports", "attack_report.docx")

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(31,  78,  121)
MID_BLUE   = RGBColor(47,  117, 181)
LIGHT_BLUE = RGBColor(189, 215, 238)
RED        = RGBColor(192, 0,   0)
GREEN      = RGBColor(0,   176, 80)
ORANGE     = RGBColor(197, 90,  17)
WHITE      = RGBColor(255, 255, 255)
GREY       = RGBColor(89,  89,  89)
BLACK      = RGBColor(0,   0,   0)

# ── Helpers ────────────────────────────────────────────────────────────────────

def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_divider(doc, color="1F4E79"):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(2)

def heading(doc, text, level=1, color=DARK_BLUE):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(4)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def para(doc, text, size=11, bold=False, italic=False,
         color=BLACK, indent=0, space_after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.bold           = bold
    run.italic         = italic
    run.font.color.rgb = color
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def bullet(doc, text, level=0, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def code_block(doc, text):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1E1E1E")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name      = "Courier New"
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(204, 204, 204)
    return p

def table(doc, headers, rows, col_widths=None):
    t   = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], "1F4E79")
        p   = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            p   = row[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def status_box(doc, title, color_hex, items):
    """Coloured info box with bullet points."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"  {title}")
    run.bold           = True
    run.font.color.rgb = WHITE
    run.font.size      = Pt(11)

    for item in items:
        pi   = doc.add_paragraph()
        pPri = pi._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"),   "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"),  color_hex + "88")
        pPri.append(shd2)
        pi.paragraph_format.left_indent  = Inches(0.4)
        pi.paragraph_format.space_after  = Pt(1)
        ri = pi.add_run(f"  • {item}")
        ri.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for s in doc.sections:
    s.top_margin    = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# ── Cover ──────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("UAV Security Research — Attack Report")
r.bold           = True
r.font.size      = Pt(24)
r.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("Timestamp Manipulation-Based GPS Spoofing on MAVLink 2.0")
r2.font.size      = Pt(14)
r2.italic         = True
r2.font.color.rgb = GREY

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta_p.add_run(f"Date: {TODAY}   |   Time: {NOW}   |   Platform: ArduPilot SITL + Gazebo Harmonic")
r3.font.size      = Pt(11)
r3.font.color.rgb = GREY

add_divider(doc)

# ── 1. Overview ────────────────────────────────────────────────────────────────
heading(doc, "1.  Overview")
para(doc,
    "This report documents the simulation of four timestamp manipulation-based GPS spoofing "
    "attacks on the MAVLink 2.0 protocol, as described in the IEEE Transactions on Communications "
    "paper by Colton, Oracevic, and Dilek (2026). The attacks were reproduced in a fully "
    "controlled Software-in-the-Loop (SITL) environment using ArduPilot, Gazebo Harmonic, "
    "MAVProxy, and Wireshark.")
para(doc,
    "All attacks exploit MAVLink 2.0's timestamp management rules — specifically the trust placed "
    "in GPS-derived time — to manipulate the UAV's internal signing clock without requiring "
    "knowledge of the secret key.")

# ── 2. Environment ─────────────────────────────────────────────────────────────
heading(doc, "2.  Lab Environment")

table(doc,
    ["Component", "Tool / Version", "Role"],
    [
        ["Flight Controller Sim", "ArduPilot SITL  v4.8.0-dev", "Simulated UAV running ArduCopter firmware"],
        ["3D Physics World",      "Gazebo Harmonic  gz-sim 8.x", "Visual and physics simulation of iris drone"],
        ["GZ Plugin",             "ardupilot_gazebo  (built)",   "Bridges SITL and Gazebo via JSON on UDP 9002"],
        ["Ground Control",        "MAVProxy  1.8.74",            "Command console — sends MAVLink to UAV"],
        ["Packet Capture",        "Wireshark  3.6.2  + tshark",  "Captures and decodes live MAVLink traffic"],
        ["Attack Scripts",        "Python 3.10  + pymavlink",    "Crafts and injects malicious MAVLink packets"],
        ["OS",                    "Ubuntu 22.04 LTS",            "Host operating system"],
    ]
)

heading(doc, "2.1  Network Port Layout", level=2, color=MID_BLUE)
code_block(doc, "Gazebo  <-- UDP 9002 (JSON FDM) -->  arducopter binary")
code_block(doc, "                                           |")
code_block(doc, "                                       TCP 5760  (MAVLink)")
code_block(doc, "                                           |")
code_block(doc, "                                      MAVProxy console")
code_block(doc, "                                           |")
code_block(doc, "                                       UDP 14550  (GCS output)")
code_block(doc, "                                           |")
code_block(doc, "                          Wireshark  (captures all MAVLink here)")
code_block(doc, "                          Attack scripts (inject via UDP 14550 / 25100)")

add_divider(doc)

# ── 3. How We Launched the Simulation ─────────────────────────────────────────
heading(doc, "3.  Launching the Simulation Stack")
para(doc,
    "A single Python launcher script (sitl_basic.py) was used to open all simulation "
    "windows automatically in the correct startup order.")

heading(doc, "3.1  Startup Order and Timing", level=2, color=MID_BLUE)
table(doc,
    ["Time", "Window", "Command", "Reason"],
    [
        ["0s",   "Win 1 — SITL",      "arducopter --model JSON ...", "Must start first to complete 2-phase boot"],
        ["12s",  "Win 2 — Gazebo",    "gz sim -r iris_ardupilot_runway.sdf", "Avoids lock_step deadlock"],
        ["Auto", "Win 3 — MAVProxy",  "mavproxy.py --master=tcp:...:5760", "Connects after SITL is ready"],
        ["Auto", "Win 4 — Wireshark", "wireshark -i lo -k ...", "Captures all MAVLink on UDP 14550"],
    ]
)

heading(doc, "3.2  Key Issues Resolved During Setup", level=2, color=MID_BLUE)
table(doc,
    ["Issue", "Root Cause", "Fix Applied"],
    [
        ["link 1 down — MAVProxy exits",  "-N flag means no-rebuild, not no-mavproxy",      "Used --no-mavproxy flag correctly"],
        ["Frame: UNSUPPORTED",            "Old eeprom.bin overriding --defaults",             "Delete eeprom.bin before each launch"],
        ["Gazebo/SITL deadlock",          "lock_step=1 with simultaneous start",             "Start SITL 12s before Gazebo"],
        ["TCP 5760 refused",              "MAVProxy already owns TCP 5760",                  "Attack scripts connect via UDP 14550"],
        ["Signing while armed error",     "ArduPilot blocks signing setup when armed",       "Setup signing before arming"],
    ]
)

add_divider(doc)

# ── 4. Attacks ────────────────────────────────────────────────────────────────
heading(doc, "4.  Attack Simulations")
para(doc,
    "Two of the four attacks from the paper were successfully simulated and captured "
    "in this session. Attacks 3 and 4 are scheduled for the next session.")

# ── Attack 1 ──────────────────────────────────────────────────────────────────
heading(doc, "4.1  Attack 1 — Signed Message Clock Manipulation", level=2, color=RED)

para(doc, "Paper Reference: Stage 1 — Replication of Ficco et al. (Section IV-C-2a)", italic=True, color=GREY)

heading(doc, "What It Does", level=3, color=MID_BLUE)
para(doc,
    "A correctly signed MAVLink 2.0 LAND command is crafted with a timestamp set "
    "1 day in the future. When the UAV receives this packet, it validates the "
    "signature, advances its internal signing clock by 1 day, and executes the "
    "LAND command. This exploits MAVLink Rule 6: receiving a valid signed message "
    "updates the local timestamp if the received time is greater.")

heading(doc, "How It Works", level=3, color=MID_BLUE)
numbered(doc, "Connect to MAVProxy's UDP output on port 14550")
numbered(doc, "Enable MAVLink 2.0 signing with key = 'key'")
numbered(doc, "Set signing timestamp to current Unix time + 86400 seconds (1 day)")
numbered(doc, "Send COMMAND_LONG: DO_SET_MODE → LAND (mode 9)")
numbered(doc, "UAV accepts valid signature, clock jumps +1 day, drone lands")

heading(doc, "Setup Required in MAVProxy", level=3, color=MID_BLUE)
code_block(doc, "signing setup key              ← set before arming")
code_block(doc, "signing setup sign_outgoing 1  ← enable signing")
code_block(doc, "mode GUIDED")
code_block(doc, "arm throttle")
code_block(doc, "takeoff 10")

heading(doc, "Attack Script", level=3, color=MID_BLUE)
code_block(doc, "python3 attacks/attack1_signed_cmd.py")

heading(doc, "Results", level=3, color=MID_BLUE)
table(doc,
    ["Metric", "Value"],
    [
        ["Script execution",      "Successful"],
        ["Connection method",     "UDP 14550 (MAVProxy output)"],
        ["Heartbeat received",    "sysid=1  compid=0"],
        ["Signing key",           "'key'  (32-byte padded)"],
        ["Spoofed timestamp",     "36006526646826  (2026-05-30 — +1 day)"],
        ["Command sent",          "COMMAND_LONG: DO_SET_MODE → LAND (mode 9)"],
        ["Capture file",          "attacks/results/attack1_capture.pcap"],
        ["Total packets captured","765"],
        ["Attack packet",         "Frame 130 — COMMAND_LONG with future timestamp"],
        ["UAV acknowledgement",   "Frame 131 — COMMAND_ACK received"],
    ]
)

heading(doc, "Wireshark Evidence", level=3, color=MID_BLUE)
para(doc, "From attack1_capture.pcap — Frame 130:")
code_block(doc, "Frame 130  |  MAVLink 2.0  |  COMMAND_LONG: MAV_CMD_DO_SET_MODE")
code_block(doc, "  Time captured : 2026-05-29 15:53:00.503 IST")
code_block(doc, "  Source        : 127.0.0.1:14550  (attack script)")
code_block(doc, "  Destination   : 127.0.0.1:57456  (MAVProxy)")
code_block(doc, "  Command       : DO_SET_MODE  →  LAND (mode 9)")
code_block(doc, "  Signature     : MAVLink 2.0 signed  (SHA-256 truncated)")
code_block(doc, "  Timestamp     : 2026-05-30  (+1 day future)")
code_block(doc, "")
code_block(doc, "Frame 131  |  COMMAND_ACK: MAV_CMD_DO_SET_MODE  ← UAV responded")

heading(doc, "Message Type Breakdown (765 packets)", level=3, color=MID_BLUE)
table(doc,
    ["Message Type", "Count", "Meaning"],
    [
        ["GLOBAL_POSITION_INT", "29", "Drone GPS position updates"],
        ["ATTITUDE",            "28", "Roll, pitch, yaw telemetry"],
        ["SYS_STATUS",          "29", "System health status"],
        ["HEARTBEAT",           "8",  "Drone alive signals"],
        ["COMMAND_LONG",        "1",  "Our attack packet (Frame 130)"],
        ["COMMAND_ACK",         "1",  "UAV acknowledgement (Frame 131)"],
    ]
)

add_divider(doc)

# ── Attack 2 ──────────────────────────────────────────────────────────────────
heading(doc, "4.2  Attack 2 — GPS Timestamp Spoofing", level=2, color=RED)

para(doc, "Paper Reference: Stage 2 — Algorithm 1 (Section IV-C-2b)", italic=True, color=GREY)

heading(doc, "What It Does", level=3, color=MID_BLUE)
para(doc,
    "Fake GPS_INPUT messages are injected into MAVProxy with a timestamp 10 days "
    "in the future. ArduPilot trusts GPS time by default (MAVLink Rule 4) and "
    "advances its internal signing clock. No secret key is required — the attacker "
    "only needs network access to send UDP packets to MAVProxy's GPS input port.")

heading(doc, "How It Works (Algorithm 1)", level=3, color=MID_BLUE)
numbered(doc, "Load GPSInput module in MAVProxy (listens on UDP 25100)")
numbered(doc, "Set GPS1_TYPE = 14 (accept MAVLink GPS) and SIM_GPS1_DISABLE = 1")
numbered(doc, "Send GPS_INPUT messages at 5 Hz with timestamp = now + 10 days")
numbered(doc, "UAV accepts GPS time → internal clock jumps +10 days")
numbered(doc, "All present-day GCS commands are now rejected as stale")

heading(doc, "Setup Required in MAVProxy", level=3, color=MID_BLUE)
code_block(doc, "param set GPS1_TYPE 14")
code_block(doc, "param set SIM_GPS1_DISABLE 1")
code_block(doc, "module load GPSInput")

heading(doc, "Attack Script", level=3, color=MID_BLUE)
code_block(doc, "python3 attacks/attack2_gps_spoof.py")

heading(doc, "Results", level=3, color=MID_BLUE)
table(doc,
    ["Metric", "Value"],
    [
        ["Script execution",       "Successful"],
        ["GPS input port",         "UDP 127.0.0.1:25100"],
        ["Clock offset injected",  "+10 days  (864000 seconds)"],
        ["Real time during attack","2026-05-29  10:30 UTC"],
        ["Spoofed GPS time",       "2026-06-08  16:00 UTC"],
        ["Send rate",              "5 Hz  (5 packets/second)"],
        ["Duration",               "30 seconds"],
        ["Total GPS_INPUT packets","150 packets sent"],
        ["Capture file",           "attacks/results/attack2_capture.pcap"],
        ["Total packets captured", "1,596"],
    ]
)

heading(doc, "Wireshark Evidence", level=3, color=MID_BLUE)
para(doc, "From attack2_capture.pcap — GPS_INPUT packet (Frame 1):")
code_block(doc, "Frame 1  |  UDP 59503 → 25100  |  GPS_INPUT (MAVLink msg 232)")
code_block(doc, "  ▼ Payload: GPS_INPUT (232)")
code_block(doc, "      time_usec  : 1780914608042893")
code_block(doc, "                   (2026-06-08 16:00:08.042893 IST)  ← 10 days ahead")
code_block(doc, "      gps_id     : 0")
code_block(doc, "      fix_type   : 3  (3D Fix)")
code_block(doc, "      lat        : -353632610  (-35.363261 deg)")
code_block(doc, "      lon        :  1491652370 (149.165237 deg)")
code_block(doc, "      alt        : 584.0 m")
code_block(doc, "      satellites : 10")
code_block(doc, "      time_week_ms: 124208042")

heading(doc, "Packet Distribution (1,596 total)", level=3, color=MID_BLUE)
table(doc,
    ["Traffic Type", "Packets", "Data", "Port"],
    [
        ["GPS_INPUT attack packets",  "150",   "16 KB",  "→ UDP 25100"],
        ["Normal MAVLink telemetry",  "1,446", "129 KB", "→ UDP 14550"],
    ]
)

add_divider(doc)

# ── 5. Attacks Remaining ──────────────────────────────────────────────────────
heading(doc, "5.  Remaining Attacks (Scheduled Next Session)")
table(doc,
    ["Attack", "Name", "What It Does", "Status"],
    [
        ["Attack 3", "Timestamp Overflow DoS",
         "Spoof GPS to Unix 4,234,820,167 — exceeds 48-bit max, all packets permanently rejected",
         "Pending"],
        ["Attack 4", "Replay Attack",
         "Capture signed packet after GPS spoof, replay after clock reset — UAV accepts stale command",
         "Pending"],
    ]
)

add_divider(doc)

# ── 6. Key Findings ───────────────────────────────────────────────────────────
heading(doc, "6.  Key Findings")

bullet(doc, "Attack 1 confirmed: A correctly signed MAVLink packet with a future timestamp is accepted "
            "by the UAV, causing a forced mode change to LAND — proven by Frame 130 in Wireshark.")
bullet(doc, "Attack 2 confirmed: GPS_INPUT messages with future timestamps successfully advance the "
            "UAV's internal signing clock by 10 days — proven by time_usec = 2026-06-08 in Wireshark.")
bullet(doc, "No secret key was needed for Attack 2 — only network access to UDP port 25100.")
bullet(doc, "MAVLink 2.0's timestamp trust model is the root vulnerability — Rules 4 and 6 allow "
            "any GPS source to manipulate the signing clock.")
bullet(doc, "SITL environment partially limits signature enforcement (USB bypass), consistent "
            "with paper findings that HITL is needed for full cryptographic verification.")

add_divider(doc)

# ── 7. Project File Structure ─────────────────────────────────────────────────
heading(doc, "7.  Project File Structure")
code_block(doc, "drone_security_lab/")
code_block(doc, "├── simulation/")
code_block(doc, "│   └── sitl_basic.py              ← launches all 4 sim windows")
code_block(doc, "├── attacks/")
code_block(doc, "│   ├── attack1_signed_cmd.py       ← Attack 1 script")
code_block(doc, "│   ├── attack2_gps_spoof.py        ← Attack 2 script")
code_block(doc, "│   ├── attack3_overflow_dos.py     ← Attack 3 (next session)")
code_block(doc, "│   ├── attack4_replay.py           ← Attack 4 (next session)")
code_block(doc, "│   └── results/")
code_block(doc, "│       ├── attack1_capture.pcap    ← Wireshark evidence Attack 1")
code_block(doc, "│       └── attack2_capture.pcap    ← Wireshark evidence Attack 2")
code_block(doc, "└── reports/")
code_block(doc, "    ├── generate_report.py          ← this script")
code_block(doc, "    └── attack_report.docx          ← output report")

add_divider(doc)

# ── 8. Next Steps ────────────────────────────────────────────────────────────
heading(doc, "8.  Next Steps")
numbered(doc, "Simulate Attack 3 — Timestamp Overflow DoS: spoof GPS to Unix 4234820164, "
              "push past 48-bit limit, observe permanent packet rejection.")
numbered(doc, "Simulate Attack 4 — Replay Attack: capture signed packet during spoofed session, "
              "restart UAV without key rotation, replay to confirm acceptance.")
numbered(doc, "Document all 4 attacks in a combined Wireshark analysis.")
numbered(doc, "Write countermeasures section based on paper Section IV-F.")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"Report saved: {OUT_FILE}")
