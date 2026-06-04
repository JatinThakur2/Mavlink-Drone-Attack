#!/usr/bin/env python3
"""Generates the May 29 daily internship report as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_FILE = os.path.join(os.path.dirname(__file__), "..", "reports", "daily_report_may29.docx")

DARK_BLUE = RGBColor(31,  78,  121)
MID_BLUE  = RGBColor(47,  117, 181)
RED       = RGBColor(192, 0,   0)
GREY      = RGBColor(89,  89,  89)
WHITE     = RGBColor(255, 255, 255)
BLACK     = RGBColor(0,   0,   0)

# ── Helpers ────────────────────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_divider(doc, color="1F4E79"):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)

def heading(doc, text, level=1, color=DARK_BLUE, space_before=10):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(2)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def para(doc, text, size=11, bold=False, italic=False,
         color=BLACK, indent=0, space_after=4):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.bold           = bold
    run.italic         = italic
    run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, indent=0.3, size=11):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def sub_bullet(doc, text):
    bullet(doc, text, indent=0.6, size=10.5)

def code(doc, text):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    return p


# ══════════════════════════════════════════════════════════════════════════════

doc = Document()
for s in doc.sections:
    s.top_margin    = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# ── Title ──────────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Internship Daily Report")
r.bold           = True
r.font.size      = Pt(20)
r.font.color.rgb = DARK_BLUE

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# MAY 29 ENTRY
# ══════════════════════════════════════════════════════════════════════════════

# Date
dp = doc.add_paragraph()
dr = dp.add_run("May 29, 2026")
dr.bold           = True
dr.font.size      = Pt(14)
dr.font.color.rgb = DARK_BLUE
dp.paragraph_format.space_after = Pt(3)

# Goal
gp = doc.add_paragraph()
gr = gp.add_run("Goal: ")
gr.bold      = True
gr.font.size = Pt(11)
gv = gp.add_run(
    "Simulate two GPS spoofing attacks on the MAVLink 2.0 protocol inside the SITL "
    "environment, capture packet evidence in Wireshark, and organise all work into a "
    "structured project folder."
)
gv.font.size = Pt(11)
gp.paragraph_format.space_after = Pt(4)

# Stack
sp = doc.add_paragraph()
sr = sp.add_run("Stack: ")
sr.bold      = True
sr.font.size = Pt(10.5)
sv = sp.add_run(
    "ArduPilot SITL v4.8.0-dev · Gazebo Harmonic (gz-sim 8.x) · ardupilot_gazebo plugin · "
    "MAVProxy 1.8.74 · Wireshark 3.6.2 · Python 3.10.12 · pymavlink 2.4.49"
)
sv.font.size      = Pt(10.5)
sv.font.color.rgb = GREY
sv.italic         = True
sp.paragraph_format.space_after = Pt(6)

add_divider(doc)

# ── Work Completed ─────────────────────────────────────────────────────────────
heading(doc, "Work Completed", level=2, color=DARK_BLUE, space_before=6)

# Project Structure
heading(doc, "Project Organisation", level=3, color=MID_BLUE, space_before=6)
bullet(doc, "Created drone_security_lab/ project folder with subfolders: simulation/, attacks/, results/, reports/.")
bullet(doc, "Moved sitl_basic.py into simulation/, created individual attack scripts for each of the 4 paper attacks.")

# Attack 1
heading(doc, "Attack 1 — Signed Message Clock Manipulation", level=3, color=MID_BLUE, space_before=6)
bullet(doc, "Replicated Ficco et al. Stage 1 from the paper (Section IV-C-2a).")
bullet(doc, "Crafted a MAVLink 2.0 LAND command signed with key = 'key' and a timestamp set +1 day in the future.")
bullet(doc, "Sent via UDP 14550 (MAVProxy output port) — TCP 5760 is owned by MAVProxy and cannot be shared.")
bullet(doc, "UAV received the signed command; Wireshark captured Frame 130 (COMMAND_LONG) and Frame 131 (COMMAND_ACK).")
bullet(doc, "765 packets captured; attack packet confirmed with future timestamp 2026-05-30 (+1 day).")

# Attack 2
heading(doc, "Attack 2 — GPS Timestamp Spoofing", level=3, color=MID_BLUE, space_before=6)
bullet(doc, "Simulated Algorithm 1 from the paper — injected fake GPS_INPUT messages with timestamp +10 days.")
bullet(doc, "Configured MAVProxy: param set GPS1_TYPE 14, SIM_GPS1_DISABLE 1, module load GPSInput.")
bullet(doc, "Sent 150 GPS_INPUT packets at 5 Hz over 30 seconds to UDP port 25100.")
bullet(doc, "Wireshark decoded time_usec = 2026-06-08 16:00:08 IST confirming the 10-day clock shift.")
bullet(doc, "1,596 total packets captured; 150 attack packets + 1,446 normal MAVLink telemetry messages.")

# Wireshark
heading(doc, "Wireshark MAVLink Dissector Fix", level=3, color=MID_BLUE, space_before=6)
bullet(doc, "Added GPS_INPUT_time_usec_fmt as a new string ProtoField in mavlink.lua.")
bullet(doc, "The new field converts raw microseconds to a human-readable date (e.g. 2026-06-08 16:00:01 UTC).")
bullet(doc, "Used as a custom Wireshark column so the spoofed date is visible directly in the packet list.")

# Reports
heading(doc, "Documentation", level=3, color=MID_BLUE, space_before=6)
bullet(doc, "Generated attack_report.docx — full technical report with methodology, Wireshark evidence, and findings.")
bullet(doc, "All pcap evidence saved: attacks/results/attack1_capture.pcap and attack2_capture.pcap.")

add_divider(doc)

# ── Issues Faced & Resolutions ────────────────────────────────────────────────
heading(doc, "Issues Faced & Resolutions", level=2, color=DARK_BLUE, space_before=6)

bullet(doc, "Signing while armed error: ArduPilot refuses signing setup when motors are armed.")
sub_bullet(doc, "Resolution: Always run 'signing setup key' before arming, not after.")

bullet(doc, "TCP 5760 connection refused: Attack script tried to connect directly to SITL, but MAVProxy already owns that port.")
sub_bullet(doc, "Resolution: Changed attack script to connect via UDP 14550 (MAVProxy's broadcast output).")

bullet(doc, "No ACK received after Attack 1: SITL bypasses signature enforcement over local USB/loopback connections.")
sub_bullet(doc, "Resolution: Expected behaviour per paper — clock manipulation effect still applied. Confirmed by Wireshark COMMAND_ACK frame.")

bullet(doc, "Wireshark column showing raw microseconds: Custom column field showed 1780914608042893 instead of a date.")
sub_bullet(doc, "Resolution: Added a new GPS_INPUT_time_usec_fmt ProtoField (string) in mavlink.lua that formats the value as a UTC date.")

bullet(doc, "gps_time_format.lua Lua error: Post-dissector tried to reference field before it was registered.")
sub_bullet(doc, "Resolution: Deleted the separate Lua file; the formatting was moved directly into mavlink.lua instead.")

add_divider(doc)

# ── Architecture ──────────────────────────────────────────────────────────────
heading(doc, "System Architecture", level=2, color=DARK_BLUE, space_before=6)
code(doc, "Gazebo  <── UDP 9002 (JSON FDM) ──>  arducopter binary")
code(doc, "                                            │")
code(doc, "                                        TCP 5760  (MAVLink)")
code(doc, "                                            │")
code(doc, "                                       MAVProxy console")
code(doc, "                                            │")
code(doc, "                                        UDP 14550  (output)")
code(doc, "                    ┌───────────────────────┤")
code(doc, "             Wireshark                 Attack scripts")
code(doc, "        (packet inspector)         (inject via UDP 14550 / 25100)")

add_divider(doc)

# ── Next Steps ────────────────────────────────────────────────────────────────
heading(doc, "Next Steps", level=2, color=DARK_BLUE, space_before=6)
bullet(doc, "Attack 3 — Timestamp Overflow DoS: spoof GPS time to Unix 4,234,820,167 (48-bit max) to permanently brick MAVLink communication.")
bullet(doc, "Attack 4 — Replay Attack: capture a signed packet during spoofed session, restart UAV without key rotation, replay to confirm acceptance.")
bullet(doc, "Document all 4 attacks with combined Wireshark analysis and add countermeasures section.")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
