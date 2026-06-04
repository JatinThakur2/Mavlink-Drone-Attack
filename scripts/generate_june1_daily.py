#!/usr/bin/env python3
"""Generates the June 1 daily internship report."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_FILE = os.path.join(os.path.dirname(__file__), "..", "reports", "daily_report_june1.docx")

DARK_BLUE = RGBColor(31,  78,  121)
MID_BLUE  = RGBColor(47,  117, 181)
GREY      = RGBColor(89,  89,  89)
WHITE     = RGBColor(255, 255, 255)
BLACK     = RGBColor(0,   0,   0)
GREEN     = RGBColor(0,   176, 80)
RED       = RGBColor(192, 0,   0)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_divider(doc, color="1F4E79"):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)

def heading(doc, text, level=1, color=DARK_BLUE, space_before=10):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(2)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def bullet(doc, text, indent=0.3, size=11):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def sub_bullet(doc, text):
    bullet(doc, text, indent=0.6, size=10.5)

def code(doc, text):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    return p

# ─────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin    = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Internship Daily Report")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK_BLUE
add_divider(doc)

# Date / Goal / Stack
dp = doc.add_paragraph()
dr = dp.add_run("June 1, 2026")
dr.bold = True; dr.font.size = Pt(14); dr.font.color.rgb = DARK_BLUE
dp.paragraph_format.space_after = Pt(3)

gp = doc.add_paragraph()
gp.add_run("Goal: ").bold = True
gp.runs[0].font.size = Pt(11)
gv = gp.add_run(
    "Implement and execute all 4 GPS spoofing attacks from the IEEE paper on the "
    "MAVLink 2.0 protocol, fix the MAVLink version to enforce 2.0 wire format (0xFD), "
    "reorganise the project folder, and capture Wireshark evidence for every attack."
)
gv.font.size = Pt(11)
gp.paragraph_format.space_after = Pt(4)

sp = doc.add_paragraph()
sp.add_run("Stack: ").bold = True
sp.runs[0].font.size = Pt(10.5)
sv = sp.add_run(
    "ArduPilot SITL v4.8.0-dev · Gazebo Harmonic · MAVProxy 1.8.74 · "
    "Wireshark 3.6.2 · Python 3.10.12 · pymavlink 2.4.49"
)
sv.font.size = Pt(10.5); sv.font.color.rgb = GREY; sv.italic = True
sp.paragraph_format.space_after = Pt(6)
add_divider(doc)

# ── Work Completed ────────────────────────────────────────────
heading(doc, "Work Completed", level=2, color=DARK_BLUE, space_before=6)

heading(doc, "Project Folder Reorganisation", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Moved mavlink-time-spoofing-main (reference code) out of attacks/ into reference/.")
bullet(doc, "Moved generate_report.py and generate_daily_report.py into scripts/ — reports/ is now output-only.")
bullet(doc, "Deleted __pycache__ and empty setup/ folder. Added README.md at project root.")

heading(doc, "MAVLink 2.0 Wire Format Fix (Attacks 1 & 2)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Both attack scripts were sending MAVLink 1.0 packets (magic byte 0xFE) — pymavlink defaults to 1.0.")
bullet(doc, "Fixed by setting os.environ[\"MAVLINK20\"] = \"1\" before importing pymavlink in both files.")
bullet(doc, "Added protocol_marker check after connecting — aborts with error if 0xFD is not confirmed.")
bullet(doc, "Re-ran both attacks; tshark confirmed 0xFD on every packet in both pcap files.")

heading(doc, "Attack 1 — Signed Message Clock Manipulation", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Signed COMMAND_LONG (DO_SET_MODE → LAND) with future timestamp (+1 day) using key = 'key'.")
bullet(doc, "Sent via UDP 14550 (MAVProxy broadcast port). MAVLink 2.0 confirmed: wire=0xFD.")
bullet(doc, "Packet timestamp: 2026-06-02 05:57:44 UTC (+1 day). Capture: attack1_capture.pcap.")

heading(doc, "Attack 2 — GPS Timestamp Spoofing", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Injected 1,253 GPS_INPUT packets at 5 Hz over 30 seconds with timestamp +10 days.")
bullet(doc, "Sent to MAVProxy GPSInput module on UDP 25100. MAVLink 2.0 confirmed: wire=0xFD.")
bullet(doc, "Capture: attack2_capture.pcap — 158 KB, 1,253 packets, 30.6 seconds.")

heading(doc, "Attack 3 — Timestamp Overflow DoS", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Injected GPS_INPUT with time_usec = 4,234,820,167,000,000 µs — the 48-bit MAVLink maximum.")
bullet(doc, "UAV signing clock forced to 2^48-1. All real 2026 packets rejected by Rule 6 (permanently).")
bullet(doc, "Verification command sent — UAV gave NO response, confirming total communication DoS.")
bullet(doc, "Capture: attack3_capture.pcap — 45 KB, 408 packets, 10.3 seconds.")

heading(doc, "Attack 4 — Replay Attack", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "SITL restarted after Attack 3 — signing clock reset to real current time.")
bullet(doc, "Signing key NOT rotated (simulating real-world operator mistake after reboot).")
bullet(doc, "Raw bytes from attack1_capture.pcap replayed byte-for-byte over UDP 14550.")
bullet(doc, "UAV responded with COMMAND_ACK (msgid=77) for command=176 (DO_SET_MODE) — replay accepted.")
bullet(doc, "Capture: attack4_capture.pcap — 47 KB, 449 packets, 14 seconds.")
add_divider(doc)

# ── Issues & Resolutions ──────────────────────────────────────
heading(doc, "Issues Faced & Resolutions", level=2, color=DARK_BLUE, space_before=6)

bullet(doc, "MAVLINK20 env var missing — both attack scripts sent MAVLink 1.0 (0xFE) packets.")
sub_bullet(doc, "Resolution: Added os.environ[\"MAVLINK20\"] = \"1\" before pymavlink import in both files.")

bullet(doc, "WIRE_PROTOCOL_VERSION attribute not found — wrong attribute name in pymavlink.")
sub_bullet(doc, "Resolution: Changed to protocol_marker (253 = 0xFD = MAVLink 2.0).")

bullet(doc, "Attack 4: conn.address returned string '0.0.0.0:14550' instead of a tuple — TypeError on port argument.")
sub_bullet(doc, "Resolution: Replaced pymavlink connection with raw UDP socket; recvfrom() returns correct (host, port) tuple.")

bullet(doc, "Pcap files corrupt — tshark used pcapng by default; SIGTERM caused incomplete final block, leaving unreadable files.")
sub_bullet(doc, "Resolution: Added -F pcap (legacy format) to all 4 attack scripts. Switched to SIGINT for clean flush on shutdown.")

bullet(doc, "Attack 4 pcap 115 MB corrupt — old corrupt file not deleted before new run; tshark overwrote the start but old data remained in the tail.")
sub_bullet(doc, "Resolution: Added os.remove(PCAP_FILE) before tshark starts. Switched to -a duration:15 so tshark auto-exits and closes the file cleanly.")

bullet(doc, "Signing setup while armed — MAVProxy rejects 'signing setup key' if motors are armed.")
sub_bullet(doc, "Resolution: Run signing setup BEFORE arming — disarm first if drone is already flying.")
add_divider(doc)

# ── Results Table ─────────────────────────────────────────────
heading(doc, "Attack Results Summary", level=2, color=DARK_BLUE, space_before=6)

tbl = doc.add_table(rows=5, cols=4)
tbl.style = "Table Grid"

headers = ["Attack", "Method", "Effect", "Pcap"]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    shade_cell(cell, "1F4E79")
    run = cell.paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(10)

rows_data = [
    ("1 — Clock Manip.", "Signed CMD +1 day ts",  "UAV clock shifted forward, LAND mode triggered",    "attack1_capture.pcap  55 KB  526 pkts"),
    ("2 — GPS Spoof",    "GPS_INPUT +10 days",    "UAV clock +10 days, GCS commands rejected",          "attack2_capture.pcap  158 KB  1,253 pkts"),
    ("3 — Overflow DoS", "GPS_INPUT at 48-bit MAX","All signed packets permanently rejected (Rule 6)",   "attack3_capture.pcap  45 KB   408 pkts"),
    ("4 — Replay",       "Raw bytes replay",       "COMMAND_ACK received — replayed command accepted",   "attack4_capture.pcap  47 KB   449 pkts"),
]
for ri, row_data in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    bg  = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row_data):
        shade_cell(row.cells[ci], bg)
        row.cells[ci].paragraphs[0].add_run(val).font.size = Pt(9.5)

doc.add_paragraph()
add_divider(doc)

# ── Next Steps ────────────────────────────────────────────────
heading(doc, "Next Steps", level=2, color=DARK_BLUE, space_before=6)
bullet(doc, "Re-run Attack 1 capture with -F pcap flag to get a clean, non-corrupt pcap file.")
bullet(doc, "Open all 4 pcap files in Wireshark and capture screenshots as visual evidence.")
bullet(doc, "Review countermeasures section and cross-reference with paper recommendations.")

doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
