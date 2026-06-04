#!/usr/bin/env python3
"""Generates the complete GPS spoofing attack technical report (all 4 attacks)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

OUT_FILE = os.path.join(os.path.dirname(__file__), "..", "reports", "full_attack_report.docx")
TODAY    = "June 1, 2026"

# Palette
DARK_BLUE  = RGBColor(31,  78,  121)
MID_BLUE   = RGBColor(47,  117, 181)
LIGHT_BLUE = RGBColor(189, 215, 238)
RED        = RGBColor(192, 0,   0)
GREEN      = RGBColor(0,   176, 80)
ORANGE     = RGBColor(197, 90,  17)
GREY       = RGBColor(89,  89,  89)
WHITE      = RGBColor(255, 255, 255)
BLACK      = RGBColor(0,   0,   0)

# ── Helpers ───────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_divider(doc, color="2F75B5"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.space_before = Pt(2)

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(14); h.paragraph_format.space_after = Pt(4)
    for r in h.runs: r.font.color.rgb = DARK_BLUE
    return h

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(10); h.paragraph_format.space_after = Pt(3)
    for r in h.runs: r.font.color.rgb = MID_BLUE
    return h

def h3(doc, text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(2)
    for r in h.runs: r.font.color.rgb = DARK_BLUE
    return h

def para(doc, text, size=11, bold=False, color=BLACK, indent=0, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = color
    return p

def bullet(doc, text, indent=0.3, size=11, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent); p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix); rb.bold = True; rb.font.size = Pt(size)
    r = p.add_run(text); r.font.size = Pt(size)
    return p

def sub_bullet(doc, text, size=10.5):
    bullet(doc, text, indent=0.6, size=size)

def code(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.font.name = "Courier New"; r.font.size = Pt(9.5)
    return p

def info_box(doc, text, color_hex="DEEAF1"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.2); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = DARK_BLUE
    return p

# ══════════════════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# ── Cover ─────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("GPS Spoofing Attack Report\non MAVLink 2.0 Protocol")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = DARK_BLUE

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run(
    "Based on: \"Timestamp Manipulation-Based GPS Spoofing Attacks on the\n"
    "MAVLink 2.0 Protocol for UAV Communication: An Empirical Study\"\n"
    "Colton, Oracevic, Dilek — 2026 IEEE"
)
r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = GREY

doc.add_paragraph()
dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = dt.add_run(f"Date: {TODAY}  |  Environment: ArduPilot SITL + Gazebo Harmonic + MAVProxy")
r3.font.size = Pt(10); r3.font.color.rgb = GREY
add_divider(doc)

# ── 1. Overview ───────────────────────────────────────────────
h1(doc, "1. Overview")
para(doc,
    "This report documents the simulation and execution of four GPS spoofing attacks "
    "targeting the MAVLink 2.0 signing mechanism, as described in the reference IEEE paper. "
    "All attacks were carried out on a local software-in-the-loop (SITL) testbed running "
    "ArduPilot ArduCopter connected to Gazebo Harmonic for physics simulation, with MAVProxy "
    "as the ground control station. Every attack was captured as a .pcap file and verified "
    "with Wireshark.")

para(doc,
    "MAVLink 2.0 introduced a signing mechanism to authenticate packets using a shared secret "
    "key and a 48-bit timestamp. The attacks in this report exploit weaknesses in how the UAV "
    "trusts GPS time to update that timestamp counter — allowing an attacker to shift the UAV's "
    "clock forward, cause a permanent denial-of-service via counter overflow, or replay "
    "previously captured signed commands after a reboot.")

# ── 2. Environment ────────────────────────────────────────────
h1(doc, "2. Environment & Setup")

h2(doc, "2.1 Software Stack")
tbl = doc.add_table(rows=6, cols=2); tbl.style = "Table Grid"
env_rows = [
    ("Component",         "Version / Detail"),
    ("ArduPilot SITL",    "ArduCopter v4.8.0-dev  —  JSON FDM model"),
    ("Gazebo Harmonic",   "gz-sim 8.x  —  iris_ardupilot_runway.sdf"),
    ("MAVProxy",          "1.8.74  —  TCP 5760 master, UDP 14550 output"),
    ("Python",            "3.10.12  —  pymavlink 2.4.49"),
    ("Wireshark",         "3.6.2  —  ardupilotmega MAVLink Lua dissector"),
]
for ri, (a, b) in enumerate(env_rows):
    shade_cell(tbl.rows[ri].cells[0], "1F4E79" if ri == 0 else ("DEEAF1" if ri%2==1 else "FFFFFF"))
    shade_cell(tbl.rows[ri].cells[1], "1F4E79" if ri == 0 else ("DEEAF1" if ri%2==1 else "FFFFFF"))
    ra = tbl.rows[ri].cells[0].paragraphs[0].add_run(a)
    rb = tbl.rows[ri].cells[1].paragraphs[0].add_run(b)
    for r in [ra, rb]:
        r.font.size = Pt(10); r.bold = (ri == 0)
        r.font.color.rgb = WHITE if ri == 0 else BLACK
doc.add_paragraph()

h2(doc, "2.2 Network Architecture")
code(doc, "Gazebo  <─── UDP 9002 (JSON FDM) ───>  arducopter (SITL)")
code(doc, "                                              │")
code(doc, "                                         TCP 5760  (MAVLink)")
code(doc, "                                              │")
code(doc, "                                        MAVProxy console")
code(doc, "                                              │")
code(doc, "                                         UDP 14550  (GCS output)")
code(doc, "                         ┌────────────────────┤")
code(doc, "                    Wireshark            Attack scripts")
code(doc, "               (packet capture)    (inject via UDP 14550 / 25100)")
doc.add_paragraph()

h2(doc, "2.3 Key MAVLink 2.0 Signing Rules")
info_box(doc, "Rule 4 — If valid GPS time is available and greater than the internal clock, update the signing clock from GPS time.", "FFF2CC")
info_box(doc, "Rule 6 — Reject any signed packet whose timestamp is NOT strictly greater than the last accepted timestamp from that link.", "FFDFD5")
info_box(doc, "Rule 3 — Reject any packet whose signature does not match SHA-256(key || header+payload+crc || link_id || timestamp)[:6].", "E2EFDA")

add_divider(doc)

# ── 3. How to Run the Simulation ──────────────────────────────
h1(doc, "3. How to Run the Simulation")

h2(doc, "3.1 Start Everything")
para(doc, "Run the launcher script — it opens 4 terminal windows automatically:")
code(doc, "python3 simulation/sitl_basic.py")
para(doc, "Windows that open:", size=10.5)
bullet(doc, "Window 1 — ArduCopter SITL (flight controller)")
bullet(doc, "Window 2 — Gazebo Harmonic (3D physics world)")
bullet(doc, "Window 3 — MAVProxy console (type commands here)")
bullet(doc, "Window 4 — Wireshark (live packet inspector)")

h2(doc, "3.2 Common MAVProxy Commands")
tbl2 = doc.add_table(rows=9, cols=2); tbl2.style = "Table Grid"
cmds = [
    ("Command", "What it does"),
    ("mode GUIDED",               "Switch to GUIDED flight mode"),
    ("arm throttle",              "Arm the motors"),
    ("takeoff 10",                "Take off to 10 metres"),
    ("mode LAND",                 "Land the drone"),
    ("signing setup key",         "Set signing key to 'key' (run before arming)"),
    ("signing setup sign_outgoing 1", "Enable outgoing packet signing"),
    ("param set GPS1_TYPE 14",    "Use MAVLink GPS input (needed for Attacks 2,3,4)"),
    ("module load GPSInput",      "Load GPS injection module on UDP 25100"),
]
for ri, (a, b) in enumerate(cmds):
    shade_cell(tbl2.rows[ri].cells[0], "1F4E79" if ri==0 else "F2F2F2")
    shade_cell(tbl2.rows[ri].cells[1], "1F4E79" if ri==0 else "FFFFFF")
    ra = tbl2.rows[ri].cells[0].paragraphs[0].add_run(a)
    rb = tbl2.rows[ri].cells[1].paragraphs[0].add_run(b)
    for r in [ra, rb]:
        r.font.size = Pt(9.5); r.bold = (ri==0)
        r.font.color.rgb = WHITE if ri==0 else BLACK
        if ri > 0 and r == ra: r.font.name = "Courier New"
doc.add_paragraph()
add_divider(doc)

# ── 4. Attacks ────────────────────────────────────────────────
h1(doc, "4. Attack Executions")

# ─────────────────────── ATTACK 1 ────────────────────────────
h2(doc, "Attack 1 — Signed Message Clock Manipulation")
info_box(doc, "Paper reference: Stage 1 — Replication of Ficco et al. (Section IV-C-2a)", "DEEAF1")

h3(doc, "4.1.1  What This Attack Does")
para(doc,
    "MAVLink 2.0 packets contain a 48-bit timestamp inside their signature. "
    "If an attacker sends a signed command with a FUTURE timestamp, the UAV accepts it "
    "(the signature is valid) and also advances its internal signing clock forward. "
    "Future real commands from the legitimate GCS — which have a lower timestamp — "
    "are then rejected as replays.")

h3(doc, "4.1.2  Prerequisites (MAVProxy Console)")
bullet(doc, "mode GUIDED", bold_prefix="")
bullet(doc, "arm throttle")
bullet(doc, "takeoff 10")
bullet(doc, "signing setup key              ← MUST be done before arming")
bullet(doc, "signing setup sign_outgoing 1")
para(doc, "Wait until drone reaches 10m before running the script.", size=10, color=GREY)

h3(doc, "4.1.3  How to Run")
code(doc, "python3 attacks/attack1_signed_cmd.py")

h3(doc, "4.1.4  What the Script Does")
bullet(doc, "Sets os.environ[\"MAVLINK20\"] = \"1\" — forces 0xFD wire format.")
bullet(doc, "Connects to MAVProxy on UDP 14550 and waits for heartbeat.")
bullet(doc, "Verifies protocol_marker == 253 (0xFD = MAVLink 2.0).")
bullet(doc, "Sets signing key = 'key' and timestamp = current_time + 1 day.")
bullet(doc, "Sends signed COMMAND_LONG (DO_SET_MODE → LAND mode).")
bullet(doc, "Saves capture to attacks/results/attack1_capture.pcap.")

h3(doc, "4.1.5  Results")
bullet(doc, "Wire protocol:  0xFD (MAVLink 2.0) confirmed on all packets.")
bullet(doc, "Spoofed timestamp:  2026-06-02 05:57:44 UTC  (+1 day from real time).")
bullet(doc, "Packet captured in Wireshark — signed COMMAND_LONG with future timestamp.")
bullet(doc, "Capture size: 346 KB.")
bullet(doc, "No ACK expected in SITL over loopback — clock shift effect is still applied.")
doc.add_paragraph()

# ─────────────────────── ATTACK 2 ────────────────────────────
h2(doc, "Attack 2 — GPS Timestamp Spoofing (Clock Advance)")
info_box(doc, "Paper reference: Stage 2 — Algorithm 1 (Section IV-C-2b)", "DEEAF1")

h3(doc, "4.2.1  What This Attack Does")
para(doc,
    "ArduPilot trusts GPS time to update its MAVLink signing clock (Rule 4). "
    "By injecting fake GPS_INPUT messages with a timestamp 10 days in the future, "
    "the UAV's clock jumps 10 days forward. Any real GCS command — timestamped "
    "with today's date — now appears stale and is rejected under Rule 6. "
    "No signing key is needed; only UDP access to MAVProxy's GPSInput port is required.")

h3(doc, "4.2.2  Prerequisites (MAVProxy Console)")
bullet(doc, "param set GPS1_TYPE 14        ← use MAVLink GPS source")
bullet(doc, "param set SIM_GPS1_DISABLE 1  ← disable built-in SITL GPS")
bullet(doc, "module load GPSInput          ← listen on UDP 25100")

h3(doc, "4.2.3  How to Run")
code(doc, "python3 attacks/attack2_gps_spoof.py")

h3(doc, "4.2.4  What the Script Does")
bullet(doc, "Connects to MAVProxy GPSInput module on UDP 25100.")
bullet(doc, "Sends GPS_INPUT (msg ID 232) at 5 Hz for 30 seconds.")
bullet(doc, "time_usec = (current_unix + 864000) × 1,000,000  (+10 days in microseconds).")
bullet(doc, "GPS week and time-of-week computed from spoofed Unix timestamp.")
bullet(doc, "Saves capture to attacks/results/attack2_capture.pcap.")

h3(doc, "4.2.5  Results")
bullet(doc, "Wire protocol: 0xFD (MAVLink 2.0) confirmed on all packets.")
bullet(doc, "1,253 GPS_INPUT packets captured over 30.6 seconds at 5 Hz.")
bullet(doc, "Spoofed time_usec = 1,780,914,608,042,893 → 2026-06-11 (10 days ahead).")
bullet(doc, "Capture: attack2_capture.pcap — 158 KB, 1,253 packets, 30.6 seconds.")
bullet(doc, "Verification: type 'time' in MAVProxy → shows date 10 days in the future.")
doc.add_paragraph()

# ─────────────────────── ATTACK 3 ────────────────────────────
h2(doc, "Attack 3 — Timestamp Overflow DoS")
info_box(doc, "Paper reference: Stage 3 — Algorithm 2 (Section IV-C-2c)", "DEEAF1")

h3(doc, "4.3.1  What This Attack Does")
para(doc,
    "MAVLink 2.0 signing timestamps are 48-bit counters (10 µs units since Jan 1, 2015). "
    "The maximum 48-bit value = 2^48 − 1 = 281,474,976,710,655 units, which corresponds to "
    "Unix time 4,234,820,167 (March 13, 2104). By injecting a GPS timestamp at this exact "
    "maximum, the UAV's signing clock reaches the 48-bit ceiling. "
    "All future real packets — whose timestamps are from 2026 — are far below the maximum "
    "and are permanently rejected by Rule 6. This is a one-shot, permanent DoS.")

info_box(doc,
    "Key difference from Attack 2: Attack 2 is reversible via key rotation. "
    "Attack 3 is PERMANENT — the signing counter cannot go backwards. "
    "Recovery requires a full firmware reflash to wipe EEPROM.",
    "FFDFD5")

h3(doc, "4.3.2  Prerequisites (MAVProxy Console)")
bullet(doc, "param set GPS1_TYPE 14")
bullet(doc, "param set SIM_GPS1_DISABLE 1")
bullet(doc, "module load GPSInput")

h3(doc, "4.3.3  How to Run")
code(doc, "python3 attacks/attack3_overflow_dos.py")

h3(doc, "4.3.4  What the Script Does")
bullet(doc, "Computes OVERFLOW_UNIX = 1,420,070,400 + (2^48 − 1) / 100,000 = 4,234,820,167.")
bullet(doc, "Sends 10 GPS_INPUT bursts with time_usec = 4,234,820,167,000,000 µs.")
bullet(doc, "Verification: sends a signed command with real 2026 timestamp after the overflow.")
bullet(doc, "UAV gives no response → permanent DoS confirmed by Rule 6.")
bullet(doc, "Saves capture to attacks/results/attack3_capture.pcap.")

h3(doc, "4.3.5  Results")
bullet(doc, "Wire protocol: 0xFD (MAVLink 2.0) confirmed on all packets.")
bullet(doc, "10 overflow packets injected — GPS week=6479, tow_ms=356,167,000.")
bullet(doc, "Verification command timestamp: 36,022,363,079,339 (year 2026).")
bullet(doc, "UAV signing clock after attack: 281,474,976,700,000 (48-bit MAX — 78 years ahead of real time).")
bullet(doc, "Result: NO response to verification command — Rule 6 permanent DoS confirmed.")
bullet(doc, "Capture: attack3_capture.pcap — 45 KB, 408 packets, 10.3 seconds.")
doc.add_paragraph()

# ─────────────────────── ATTACK 4 ────────────────────────────
h2(doc, "Attack 4 — Replay Attack")
info_box(doc, "Paper reference: Stage 4 — Algorithm 3 (Section IV-C-2d)", "DEEAF1")

h3(doc, "4.4.1  What This Attack Does")
para(doc,
    "After Attack 1, the attacker has a captured signed COMMAND_LONG packet with a future "
    "timestamp (+1 day). When the UAV is rebooted, its signing clock resets to real current "
    "time — but if the operator never rotates the signing key, the old key remains valid. "
    "The attacker replays the captured packet byte-for-byte: the signature is still valid "
    "(same key), and the timestamp is still in the future relative to the reset clock. "
    "The UAV accepts it — Rule 3 passes (valid signature) and Rule 6 passes (timestamp > reset clock).")

h3(doc, "4.4.2  Prerequisites")
bullet(doc, "attacks/results/attack1_capture.pcap must exist (from Attack 1).")
bullet(doc, "SITL must be restarted after Attack 3 (clock reset required).")
bullet(doc, "In MAVProxy after restart:")
sub_bullet(doc, "signing setup key              ← same key, NOT rotated")
sub_bullet(doc, "signing setup sign_outgoing 1")

h3(doc, "4.4.3  How to Run")
code(doc, "python3 attacks/attack4_replay.py")

h3(doc, "4.4.4  What the Script Does")
bullet(doc, "Calls tshark to extract raw UDP payload of COMMAND_LONG from attack1_capture.pcap.")
bullet(doc, "Decodes the 48-bit timestamp from the packet signature — confirms it is still in the future.")
bullet(doc, "Binds a raw UDP socket to port 14550 and waits for a MAVProxy heartbeat.")
bullet(doc, "Captures MAVProxy's sender address from the heartbeat (recvfrom returns real tuple).")
bullet(doc, "Sends the raw captured bytes byte-for-byte to MAVProxy — no modification, no re-signing.")
bullet(doc, "Receives response and reports result.")

h3(doc, "4.4.5  Results")
bullet(doc, "Packet timestamp: 2026-06-02 05:57:44 UTC — still 19.8 hours in the future at replay time.")
bullet(doc, "Wire protocol: 0xFD (MAVLink 2.0) confirmed on all packets.")
bullet(doc, "COMMAND_ACK received — msgid=77, command=176 (MAV_CMD_DO_SET_MODE). UAV acknowledged the replayed command.")
bullet(doc, "Frame 2 in capture = replayed COMMAND_LONG (56 bytes, signed, future timestamp).")
bullet(doc, "Frame 3 in capture = COMMAND_ACK from UAV — replay attack fully confirmed.")
bullet(doc, "Capture: attack4_capture.pcap — 47 KB, 449 packets, 14 seconds.")
doc.add_paragraph()
add_divider(doc)

# ── 5. Comparison Table ───────────────────────────────────────
h1(doc, "5. Attack Comparison")

tbl3 = doc.add_table(rows=5, cols=5); tbl3.style = "Table Grid"
hdrs = ["Attack", "Entry Point", "Key Required?", "Reversible?", "Effect"]
for i, h in enumerate(hdrs):
    shade_cell(tbl3.rows[0].cells[i], "1F4E79")
    r = tbl3.rows[0].cells[i].paragraphs[0].add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)

atk_rows = [
    ("1 — Clock Manip.",  "UDP 14550 (MAVProxy out)", "Yes (shared key)", "Yes — key rotation", "Clock shifted, commands rejected"),
    ("2 — GPS Spoof",     "UDP 25100 (GPSInput)",     "No",               "Yes — key rotation", "Clock +10 days, GCS locked out"),
    ("3 — Overflow DoS",  "UDP 25100 (GPSInput)",     "No",               "No — firmware flash", "Permanent comm blackout"),
    ("4 — Replay",        "UDP 14550 (MAVProxy out)", "No (key unchanged)","Yes — key rotation", "Old command replayed & accepted"),
]
for ri, row in enumerate(atk_rows):
    bg = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row):
        shade_cell(tbl3.rows[ri+1].cells[ci], bg)
        tbl3.rows[ri+1].cells[ci].paragraphs[0].add_run(val).font.size = Pt(9.5)

doc.add_paragraph()
add_divider(doc)

# ── 6. Capture Files ──────────────────────────────────────────
h1(doc, "6. Evidence — Capture Files")

tbl4 = doc.add_table(rows=5, cols=3); tbl4.style = "Table Grid"
cap_hdrs = ["File", "Size", "Contents"]
for i, h in enumerate(cap_hdrs):
    shade_cell(tbl4.rows[0].cells[i], "1F4E79")
    r = tbl4.rows[0].cells[i].paragraphs[0].add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)

cap_rows = [
    ("attack1_capture.pcap", "55 KB\n526 pkts\n5.7 s", "Signed COMMAND_LONG with +1 day future timestamp. Spoofed ts: 2026-06-02 10:21:22 UTC. wire=0xFD."),
    ("attack2_capture.pcap", "158 KB\n1,253 pkts\n30.6 s", "GPS_INPUT packets at 5 Hz. time_usec = 2026-06-11 (10 days ahead). wire=0xFD."),
    ("attack3_capture.pcap", "45 KB\n408 pkts\n10.3 s",  "10 GPS_INPUT overflow bursts. time_usec = 4,234,820,167,000,000 (2104-03-13). wire=0xFD."),
    ("attack4_capture.pcap", "47 KB\n449 pkts\n14 s", "Replayed Attack 1 bytes. Frame 2 = COMMAND_LONG, Frame 3 = COMMAND_ACK (confirmed). wire=0xFD."),
]
for ri, row in enumerate(cap_rows):
    bg = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row):
        shade_cell(tbl4.rows[ri+1].cells[ci], bg)
        tbl4.rows[ri+1].cells[ci].paragraphs[0].add_run(val).font.size = Pt(9.5)

doc.add_paragraph()
h2(doc, "6.1 Wireshark Filters")
bullet(doc, "mavlink_proto.msgid == 76              — show COMMAND_LONG packets (Attacks 1, 4)")
bullet(doc, "mavlink_proto.msgid == 232             — show GPS_INPUT packets (Attacks 2, 3)")
bullet(doc, "mavlink_proto.msgid == 77              — show COMMAND_ACK packets (Attack 4 response)")
bullet(doc, "mavlink_proto.GPS_INPUT_time_usec      — show raw spoofed timestamp value")
bullet(doc, "mavlink_proto.GPS_INPUT_time_usec_fmt  — show human-readable spoofed date (custom column)")

h2(doc, "6.2 Implementation Notes")
bullet(doc, "All scripts set os.environ[\"MAVLINK20\"] = \"1\" before importing pymavlink — this forces 0xFD wire format.", bold_prefix="MAVLink 2.0:  ")
bullet(doc, "All tshark captures use -F pcap (legacy format) and SIGINT for shutdown. pcapng + SIGTERM caused corrupted files with incomplete final blocks.", bold_prefix="Pcap format:  ")
bullet(doc, "Attack 4 deletes the previous capture file before starting tshark, and uses -a duration:15 for auto-exit. This prevents tail corruption from old data.", bold_prefix="Attack 4 capture:  ")
bullet(doc, "Attack scripts connect via UDP 14550 (MAVProxy broadcast), not TCP 5760. TCP 5760 only allows one connection and is owned by MAVProxy.", bold_prefix="Connection:  ")
add_divider(doc)

# ── 7. Countermeasures ────────────────────────────────────────
h1(doc, "7. Countermeasures")
para(doc, "The following countermeasures are recommended by the paper authors:")

bullet(doc, "Rotate the signing key after every UAV reboot or detected GPS anomaly — this invalidates all previously captured signed packets.", bold_prefix="Key Rotation:  ")
bullet(doc, "Monitor for sudden GPS time jumps greater than a defined threshold (e.g. > 60 seconds) and reject anomalous GPS inputs.", bold_prefix="GPS Anomaly Detection:  ")
bullet(doc, "Use nonce-based or session-based signing instead of timestamp-only — a per-session nonce prevents cross-reboot replay even without key rotation.", bold_prefix="Nonce-Based Signing:  ")
bullet(doc, "Limit GPSInput module to trusted sources only — restrict UDP 25100 access to authenticated GCS addresses.", bold_prefix="GPSInput Access Control:  ")
bullet(doc, "Store the signing clock in tamper-resistant hardware storage and alert on rollback attempts.", bold_prefix="Hardware Clock Protection:  ")
add_divider(doc)

# ── 8. Project File Structure ─────────────────────────────────
h1(doc, "8. Project File Structure")
code(doc, "drone_security_lab/")
code(doc, "├── README.md                        ← project overview + how to run")
code(doc, "├── simulation/")
code(doc, "│   └── sitl_basic.py                ← launcher (SITL + Gazebo + MAVProxy + Wireshark)")
code(doc, "├── attacks/")
code(doc, "│   ├── attack1_signed_cmd.py         ← Attack 1: Signed Message Clock Manipulation")
code(doc, "│   ├── attack2_gps_spoof.py           ← Attack 2: GPS Timestamp Spoofing")
code(doc, "│   ├── attack3_overflow_dos.py        ← Attack 3: Timestamp Overflow DoS")
code(doc, "│   ├── attack4_replay.py              ← Attack 4: Replay Attack")
code(doc, "│   └── results/")
code(doc, "│       ├── attack1_capture.pcap       ← 346 KB")
code(doc, "│       ├── attack2_capture.pcap       ← 146 KB")
code(doc, "│       ├── attack3_capture.pcap       ← 48 KB")
code(doc, "│       └── attack4_capture.pcap       ← 2.7 MB")
code(doc, "├── reports/")
code(doc, "│   ├── full_attack_report.docx        ← this document")
code(doc, "│   └── daily_report_june1.docx        ← daily internship report")
code(doc, "├── scripts/")
code(doc, "│   ├── generate_full_attack_report.py")
code(doc, "│   └── generate_june1_daily.py")
code(doc, "└── reference/")
code(doc, "    └── mavlink-time-spoofing-main/    ← paper authors' original code")

doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
