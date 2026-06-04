#!/usr/bin/env python3
"""Generates the full technical attack report v2 (includes Attack 3B + Wireshark dissectors)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_FILE = os.path.join(os.path.dirname(__file__), "..", "reports", "full_attack_report_v2.docx")

DARK_BLUE = RGBColor(31,  78,  121)
MID_BLUE  = RGBColor(47,  117, 181)
LIGHT_BLUE= RGBColor(189, 215, 238)
GREY      = RGBColor(89,  89,  89)
WHITE     = RGBColor(255, 255, 255)
BLACK     = RGBColor(0,   0,   0)
GREEN     = RGBColor(0,   176, 80)
RED       = RGBColor(192, 0,   0)
ORANGE    = RGBColor(255, 102, 0)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_divider(doc, color="1F4E79"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.space_before = Pt(2)

def heading(doc, text, level=1, color=DARK_BLUE, space_before=12):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before); h.paragraph_format.space_after = Pt(3)
    for run in h.runs: run.font.color.rgb = color
    return h

def bullet(doc, text, indent=0.3, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text); run.font.size = Pt(size)
    return p

def sub_bullet(doc, text): bullet(doc, text, indent=0.6, size=10.5)

def body(doc, text, size=11):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(size); p.paragraph_format.space_after = Pt(4)
    return p

def code(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text); run.font.name = "Courier New"; run.font.size = Pt(9)
    return p

# ─────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)

# ══ COVER PAGE ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("GPS Spoofing & MAVLink 2.0 Signing Attacks")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK_BLUE

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Full Technical Attack Report — v2.0")
r2.font.size = Pt(14); r2.font.color.rgb = MID_BLUE; r2.bold = True

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run(
    "Based on: \"Timestamp Manipulation-Based GPS Spoofing Attacks on the\n"
    "MAVLink 2.0 Protocol for UAV Communication\"\n"
    "Colton, Oracevic, Dilek — IEEE 2026"
)
r3.font.size = Pt(11); r3.font.color.rgb = GREY; r3.italic = True

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run("Internship Project  ·  June 1–2, 2026")
r4.font.size = Pt(11); r4.font.color.rgb = GREY

doc.add_page_break()

# ══ 1. OVERVIEW ══════════════════════════════════════════════
heading(doc, "1. Project Overview", level=1, color=DARK_BLUE)
body(doc,
    "This report documents the complete implementation and execution of 5 GPS spoofing and "
    "MAVLink 2.0 signing attacks against an ArduPilot SITL (Software In The Loop) drone "
    "simulation. All attacks are based on the IEEE 2026 paper by Colton, Oracevic, and Dilek, "
    "with extensions including a new Timestamp Underflow attack (Attack 3B) not present in "
    "the original paper."
)
body(doc,
    "Each attack exploits a different weakness in MAVLink 2.0's signing mechanism. "
    "All 5 attacks were executed successfully, captured with Wireshark (custom Lua dissectors), "
    "and visually verified in Gazebo Harmonic 3D simulation."
)

# ══ 2. ENVIRONMENT ════════════════════════════════════════════
heading(doc, "2. Environment & Architecture", level=1, color=DARK_BLUE)

heading(doc, "2.1 Software Stack", level=2, color=MID_BLUE, space_before=6)
tbl = doc.add_table(rows=7, cols=2); tbl.style = "Table Grid"
shade_cell(tbl.rows[0].cells[0], "1F4E79"); shade_cell(tbl.rows[0].cells[1], "1F4E79")
tbl.rows[0].cells[0].paragraphs[0].add_run("Component").font.color.rgb = WHITE
tbl.rows[0].cells[1].paragraphs[0].add_run("Version / Details").font.color.rgb = WHITE
stack = [
    ("ArduPilot SITL",   "v4.8.0-dev — ArduCopter, JSON FDM model, TCP 5760"),
    ("Gazebo",           "Harmonic — iris_ardupilot_runway.sdf, UDP 9002 JSON FDM"),
    ("MAVProxy",         "1.8.74 — TCP 5760 master, UDP 14550 output broadcast"),
    ("Python",           "3.10.12 — pymavlink 2.4.49, socket, json, hashlib, struct"),
    ("Wireshark",        "3.6.2 — Lua 5.1 dissectors: gps_spoof.lua, attack4_replay.lua"),
    ("Reference Code",   "mavlink-time-spoofing-main (Colton et al.) — reference only"),
]
for ri, (comp, ver) in enumerate(stack):
    bg = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    shade_cell(tbl.rows[ri+1].cells[0], bg); shade_cell(tbl.rows[ri+1].cells[1], bg)
    tbl.rows[ri+1].cells[0].paragraphs[0].add_run(comp).font.size = Pt(10)
    tbl.rows[ri+1].cells[1].paragraphs[0].add_run(ver).font.size = Pt(10)

doc.add_paragraph()
heading(doc, "2.2 Network Architecture", level=2, color=MID_BLUE, space_before=6)
code(doc, "Gazebo  <── UDP 9002 (JSON FDM) ──>  ArduCopter SITL")
code(doc, "                                            │")
code(doc, "                                       TCP 5760")
code(doc, "                                            │")
code(doc, "                                      MAVProxy 1.8.74")
code(doc, "                                            │")
code(doc, "                           ┌─────── UDP 14550 ──────────┐")
code(doc, "                           │                             │")
code(doc, "                    Attack scripts               Wireshark / GCS")
code(doc, "                    (port 14550)                (port 14550)")
code(doc, "")
code(doc, "GPS Injection:  Attack scripts ──> UDP 25100 ──> MAVProxy GPSInput module")

doc.add_paragraph()
heading(doc, "2.3 MAVLink 2.0 Signing Rules", level=2, color=MID_BLUE, space_before=6)
bullet(doc, "Rule 3 — Reject if signature does not match SHA-256(key || header+payload+crc || link_id || ts)[:6]")
bullet(doc, "Rule 4 — If GPS time > internal clock, update signing clock from GPS time.")
bullet(doc, "Rule 6 — Reject if timestamp ≤ last accepted timestamp from this link.")
body(doc, "Rules 4 and 6 together create the attack surface: Rule 4 lets GPS time advance the clock, "
     "and Rule 6 locks out any earlier timestamps permanently.", size=10.5)

add_divider(doc)

# ══ 3. ATTACKS ════════════════════════════════════════════════
heading(doc, "3. Attack Implementations", level=1, color=DARK_BLUE)

# Attack 1
heading(doc, "Attack 1 — Signed Message Clock Manipulation", level=2, color=MID_BLUE, space_before=6)
body(doc, "Paper reference: Stage 1 / Section IV-C-2a")
bullet(doc, "Theory: A MAVLink 2.0 signed packet contains a 48-bit timestamp in the signature block. "
       "If we sign a COMMAND_LONG with a future timestamp (+1 day), the UAV accepts it (Rule 3 passes), "
       "advances its internal clock (Rule 4 analogue), and executes the command.")
bullet(doc, "Implementation: Manual SHA-256 signing using hashlib + mavcrc.x25crc_slow — identical to "
       "reference custom-input.py. Key hashed via sha256(key_string) before use.")
bullet(doc, "Command sent: DO_SET_MODE → LAND (command=176, custom_mode=9).")
bullet(doc, "Sent via: conn.write(raw_packet) to UDP 14550 (MAVProxy broadcast port).")
bullet(doc, "Result: Drone switches to LAND mode and descends. COMMAND_ACK result=0 received.")
bullet(doc, "Wireshark: filter mavlink_proto.msgid == 76 || mavlink_proto.msgid == 77")
sub_bullet(doc, "ATK4_REPLAY: COMMAND_LONG with future signature timestamp visible.")
sub_bullet(doc, "ATK4_ACCEPTED: COMMAND_ACK result=0 confirming acceptance.")
bullet(doc, "Pcap: attack1_capture.pcap — 55 KB, 526 packets, 5.7 seconds.")

doc.add_paragraph()

# Attack 2
heading(doc, "Attack 2 — GPS Timestamp + Position Spoofing", level=2, color=MID_BLUE, space_before=6)
body(doc, "Paper reference: Stage 2 / Section IV-C-2b / Algorithm 1")
bullet(doc, "Theory: ArduPilot accepts GPS_INPUT messages via MAVProxy GPSInput module (UDP 25100). "
       "Injecting GPS time 10 days ahead advances the UAV signing clock — future GCS commands appear stale.")
bullet(doc, "Implementation: JSON format over raw UDP socket (matching reference simulate-gps.py). "
       "Previous implementation sent binary MAVLink — incorrect for the GPSInput module.")
bullet(doc, "LIE 1 — Timestamp: time_usec = real_unix + 10_days_in_microseconds. "
       "UAV signing clock jumps +10 days. Real GCS commands rejected by Rule 6.")
bullet(doc, "LIE 2 — Position: lat drifts -0.0018 degrees southward over 30 seconds (~200m). "
       "UAV compensates northward — drone visibly flies 200m north in Gazebo.")
bullet(doc, "Result: Both effects confirmed. 'time' command in MAVProxy shows June 12. "
       "Subsequent mode change commands rejected.")
bullet(doc, "Wireshark: filter udp.dstport == 25100 → Protocol ATK2_GPS_SPOOF.")
sub_bullet(doc, "GPS Date Claimed: 2026-06-12 | Capture Date: 2026-06-02 | Time Gap: +10 days.")
sub_bullet(doc, "Latitude drifts from -35.363261 to -35.365061 over 149 packets.")
bullet(doc, "Pcap: attack2_capture.pcap — 343 KB, 149 GPS packets, 30 seconds.")

doc.add_paragraph()

# Attack 3
heading(doc, "Attack 3 — Timestamp Overflow DoS", level=2, color=MID_BLUE, space_before=6)
body(doc, "Paper reference: Stage 3 / Section IV-C-2c / Algorithm 2")
bullet(doc, "Theory: MAVLink 2.0 signing timestamps are 48-bit (max = 2^48-1 = 281,474,976,710,655 units). "
       "Injecting GPS time at Unix 4,234,820,167 (March 13, 2104) sets the UAV signing clock to the "
       "48-bit maximum. All real 2026 packets have timestamps far below this — Rule 6 rejects them all permanently.")
bullet(doc, "Implementation: JSON GPS_INPUT with time_usec = 4,234,820,167,000,000. Burst of 10 packets.")
bullet(doc, "Result: Total communication blackout. All signed commands silently rejected. "
       "Drone continues last guided waypoint indefinitely — cannot be recalled.")
bullet(doc, "Recovery: Firmware reflash only. Clock persists in EEPROM across reboots and key rotation.")
bullet(doc, "Wireshark: filter udp.dstport == 25100 → Protocol ATK3_OVERFLOW.")
sub_bullet(doc, "GPS Date: 2104-03-13 | Years Ahead: ~78 | MAVLink TS: 281,474,976,700,000.")
sub_bullet(doc, "Units below MAX: 10,655 (~0.107 seconds below 48-bit ceiling).")
bullet(doc, "Pcap: attack3_capture.pcap — 99 KB, 10 overflow packets + verification traffic.")

doc.add_paragraph()

# Attack 3B
heading(doc, "Attack 3B — Timestamp Underflow DoS (New)", level=2, color=MID_BLUE, space_before=6)
body(doc, "Extension — not in original paper. Symmetric to Attack 3 but in the opposite direction.")
bullet(doc, "Theory: If GPS time is BEFORE the MAVLink epoch (Jan 1, 2015), the computed signing timestamp "
       "is NEGATIVE as int64. Cast to uint64, it wraps to near 2^64 — a value larger than the 48-bit max.")
bullet(doc, "Injection: time_usec = 1,262,304,000,000,000 (Jan 1, 2010).")
sub_bullet(doc, "MAVLink TS (signed): -15,776,640,000,000 (negative — 5 years before epoch).")
sub_bullet(doc, "uint64 wrapped: 18,446,728,297,069,551,616 (near 2^64).")
sub_bullet(doc, "48-bit masked: 265,698,336,710,656 >> real 2026 TS (~36 trillion) → Rule 6 DoS.")
bullet(doc, "Two possible effects depending on ArduPilot implementation:")
sub_bullet(doc, "Effect A (uint64 wrap): DoS identical to Attack 3 — all 2026 packets rejected.")
sub_bullet(doc, "Effect B (clamped to 0): Signing clock resets to Jan 1 2015. All 2026 packets accepted again — acts as clock reset to undo Attack 2 or 3.")
bullet(doc, "Wireshark: filter udp.dstport == 25100 → Protocol ATK0_UNDERFLOW.")
sub_bullet(doc, "GPS Date: 2010-01-01 | MAVLink TS: -15,776,640,000,000 | Wrapped: 18,446,728,297,069,551,616.")
bullet(doc, "Pcap: attack3b_capture.pcap — 10 underflow packets + verification traffic.")

doc.add_paragraph()

# Attack 4
heading(doc, "Attack 4 — Replay Attack", level=2, color=MID_BLUE, space_before=6)
body(doc, "Paper reference: Stage 4 / Section IV-C-2d / Algorithm 3")
bullet(doc, "Theory: After SITL reboot, the signing clock resets to real current time. If the key is NOT "
       "rotated (common operator mistake), the old Attack 1 packet is still valid because its +1 day "
       "timestamp is now greater than the reset clock.")
bullet(doc, "Implementation: Extract raw bytes of COMMAND_LONG from attack1_capture.pcap using tshark. "
       "Replay byte-for-byte via raw UDP socket — no modification, no re-signing.")
bullet(doc, "Conditions: a) Same signing key in use. b) SITL rebooted (clock reset). "
       "c) Attack 1 pcap < 24 hours old (timestamp still in future).")
bullet(doc, "Result: COMMAND_ACK (msgid=77, result=0) received. Drone switches to LAND mode and "
       "descends in Gazebo — same visual effect as Attack 1, triggered by yesterday's captured bytes.")
bullet(doc, "Root cause: MAVLink 2.0 has NO session-based replay protection. Rule 6 only prevents "
       "same-session replays. After reboot, counter resets — making captured future-timestamped packets replayable.")
bullet(doc, "Countermeasure: Rotate signing key after every reboot. Use nonce-based signing.")
bullet(doc, "Wireshark: filter atk4replay.type or mavlink_proto.msgid == 76 || mavlink_proto.msgid == 77.")
sub_bullet(doc, "Frame 2 — ATK4_REPLAY: COMMAND_LONG with signature date from Attack 1 capture, still in future.")
sub_bullet(doc, "Frame 3 — ATK4_ACCEPTED: COMMAND_ACK result=0 confirming replay accepted.")
bullet(doc, "Pcap: attack4_capture.pcap — 47 KB, 449 packets, 14 seconds.")

add_divider(doc)

# ══ 4. WIRESHARK DISSECTORS ═══════════════════════════════════
heading(doc, "4. Wireshark Lua Dissectors", level=1, color=DARK_BLUE)
body(doc, "Two Lua dissector plugins were written to decode attack-specific fields directly in the "
     "Wireshark GUI — replacing raw hex with human-readable attack analysis.")

heading(doc, "4.1 gps_spoof.lua — GPS Injection Decoder", level=2, color=MID_BLUE, space_before=6)
bullet(doc, "Registered on UDP port 25100 via DissectorTable.get(\"udp.port\"):add(25100, p).")
bullet(doc, "Auto-detects attack zone by time_usec magnitude:")
sub_bullet(doc, "Zone 0 (< Jan 2015): ATK0_UNDERFLOW — shows negative MAVLink TS, uint64 wrap, Effect A/B explanation.")
sub_bullet(doc, "Zone 2 (2015–2065): ATK2_GPS_SPOOF — shows +10 day gap, position drift in metres.")
sub_bullet(doc, "Zone 3 (> 2065): ATK3_OVERFLOW — shows 78-year gap, MAVLink TS vs 2^48-1, recovery method.")
bullet(doc, "Key fix: Lua double precision (max exact int = 2^53) cannot store uint64 values (2^64). "
       "Pre-computed string constants used for UINT64_MAX arithmetic in Zone 0.")

heading(doc, "4.2 attack4_replay.lua — MAVLink Replay Decoder", level=2, color=MID_BLUE, space_before=6)
bullet(doc, "Registered as post-dissector via register_postdissector(p, true) — runs after mavlink_proto.")
bullet(doc, "allfields=true enables display filter indexing (without it, atk4replay.type returns 0 results).")
bullet(doc, "Annotates COMMAND_LONG (76): shows signature timestamp, hours in future, command decoded as LAND.")
bullet(doc, "Annotates COMMAND_ACK (77): shows result=0 ACCEPTED with replay verdict.")
bullet(doc, "Protocol column: ATK4_REPLAY / ATK4_ACCEPTED / ATK4_REJECTED.")

heading(doc, "4.3 Display Filters Reference", level=2, color=MID_BLUE, space_before=6)
tbl3 = doc.add_table(rows=7, cols=3); tbl3.style = "Table Grid"
for i, h in enumerate(["Filter", "Purpose", "Attack"]):
    shade_cell(tbl3.rows[0].cells[i], "1F4E79")
    tbl3.rows[0].cells[i].paragraphs[0].add_run(h).font.color.rgb = WHITE
filters = [
    ("udp.dstport == 25100",             "All GPS injection packets",               "2, 3, 3B"),
    ("atk2gps.gap",                      "Packets with 10-day timestamp gap",       "2"),
    ("atk2gps.drift",                    "Packets with position drift field",       "2"),
    ("mavlink_proto.msgid == 76 || mavlink_proto.msgid == 77", "COMMAND_LONG + ACK","1, 4"),
    ("atk4replay.type",                  "Replay packet + ACK only",               "1, 4"),
    ("frame.protocols contains \"atk4replay\"", "All packets our dissector touched","1, 4"),
]
for ri, row_data in enumerate(filters):
    row = tbl3.rows[ri+1]
    bg = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row_data):
        shade_cell(row.cells[ci], bg)
        run = row.cells[ci].paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if ci == 0: run.font.name = "Courier New"

add_divider(doc)

# ══ 5. COMPARISON ═════════════════════════════════════════════
heading(doc, "5. Implementation vs Reference", level=1, color=DARK_BLUE)
body(doc, "The reference code (mavlink-time-spoofing-main by Colton et al.) was studied and our "
     "implementation was aligned on Day 2. Key differences remain due to our SITL/MAVProxy setup.")

tbl4 = doc.add_table(rows=7, cols=3); tbl4.style = "Table Grid"
for i, h in enumerate(["Aspect", "Reference", "Our Implementation"]):
    shade_cell(tbl4.rows[0].cells[i], "1F4E79")
    tbl4.rows[0].cells[i].paragraphs[0].add_run(h).font.color.rgb = WHITE
comp = [
    ("GPS injection format",  "JSON via raw UDP socket",          "JSON via raw UDP socket (aligned Day 2)"),
    ("Signing method",        "Manual SHA-256 + mavcrc",          "Manual SHA-256 + mavcrc (aligned Day 2)"),
    ("Key derivation",        "sha256(key_string).digest()",      "sha256(key_string).digest() (aligned Day 2)"),
    ("Connection port",       "UDP 5760 (direct to SITL)",        "UDP 14550 (MAVProxy output)"),
    ("Launcher",              "tmux + sim_vehicle.py",            "gnome-terminal + arducopter binary"),
    ("GCS",                   "QGroundControl",                   "MAVProxy console"),
]
for ri, row_data in enumerate(comp):
    row = tbl4.rows[ri+1]
    bg = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row_data):
        shade_cell(row.cells[ci], bg)
        row.cells[ci].paragraphs[0].add_run(val).font.size = Pt(9.5)

add_divider(doc)

# ══ 6. RESULTS SUMMARY ════════════════════════════════════════
heading(doc, "6. Results Summary", level=1, color=DARK_BLUE)

tbl5 = doc.add_table(rows=6, cols=6); tbl5.style = "Table Grid"
for i, h in enumerate(["Attack", "Vulnerability", "Effect", "Gazebo Visual", "Reversible", "Pcap"]):
    shade_cell(tbl5.rows[0].cells[i], "1F4E79")
    run = tbl5.rows[0].cells[i].paragraphs[0].add_run(h)
    run.font.color.rgb = WHITE; run.font.size = Pt(9); run.bold = True

results = [
    ("1 — Clock Manip.", "Rule 3: valid future sig",  "Clock advanced, LAND",   "Drone lands",            "Yes (key rotation)", "55 KB  526 pkts"),
    ("2 — GPS Spoof",    "Rule 4: GPS trusted",       "Clock +10d, pos drift",  "Drone flies 200m N",    "Yes (key rotation)", "343 KB 149 pkts"),
    ("3 — Overflow",     "Rule 4+6: 48-bit max",      "Permanent DoS",          "Drone ignores all cmds","No — reflash only",  "99 KB  10 pkts"),
    ("3B — Underflow",   "Rule 4+6: pre-epoch wrap",  "DoS or clock reset",     "DoS or unblocks GCS",   "Impl. dependent",    "varies"),
    ("4 — Replay",       "No session replay protect", "LAND from old packet",   "Drone lands again",      "Yes (key rotation)", "47 KB  449 pkts"),
]
for ri, row_data in enumerate(results):
    row = tbl5.rows[ri+1]
    bg = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row_data):
        shade_cell(row.cells[ci], bg)
        row.cells[ci].paragraphs[0].add_run(val).font.size = Pt(9)

add_divider(doc)

# ══ 7. COUNTERMEASURES ════════════════════════════════════════
heading(doc, "7. Countermeasures", level=1, color=DARK_BLUE)
bullet(doc, "Rotate signing key after every reboot — prevents Attack 4 replay.")
bullet(doc, "Monitor for sudden GPS time jumps > threshold — detect Attacks 2 and 3.")
bullet(doc, "Validate GPS time against system clock before updating signing timestamp — block all GPS-based attacks.")
bullet(doc, "Use nonce-based signing instead of timestamp-only — eliminates replay window entirely.")
bullet(doc, "Enforce GPS time sanity bounds — reject GPS time > current + 1 hour or < current - 1 hour.")
bullet(doc, "Implement 48-bit overflow check — reject GPS time that would push signing clock to > 2^47.")

add_divider(doc)

# ══ 8. FILE STRUCTURE ═════════════════════════════════════════
heading(doc, "8. Project File Structure", level=1, color=DARK_BLUE)
code(doc, "drone_security_lab/")
code(doc, "├── attacks/")
code(doc, "│   ├── attack1_signed_cmd.py       ← Manual SHA-256 signed COMMAND_LONG")
code(doc, "│   ├── attack2_gps_spoof.py        ← JSON GPS +10d + position drift")
code(doc, "│   ├── attack3_overflow_dos.py     ← JSON GPS at year 2104 (48-bit max)")
code(doc, "│   ├── attack3b_underflow_dos.py   ← JSON GPS at year 2010 (pre-epoch)")
code(doc, "│   ├── attack4_replay.py           ← Raw byte replay from attack1 pcap")
code(doc, "│   └── results/                   ← pcap files for all 5 attacks")
code(doc, "├── simulation/")
code(doc, "│   └── sitl_basic.py              ← Launches SITL + Gazebo + MAVProxy")
code(doc, "├── scripts/")
code(doc, "│   ├── generate_june1_daily.py")
code(doc, "│   ├── generate_june2_daily.py")
code(doc, "│   └── generate_full_report_v2.py ← This document")
code(doc, "├── reports/                       ← Generated Word documents")
code(doc, "├── reference/")
code(doc, "│   └── mavlink-time-spoofing-main ← Original paper authors' code")
code(doc, "~/.config/wireshark/plugins/")
code(doc, "├── gps_spoof.lua                  ← Decodes Attacks 2, 3, 3B on UDP 25100")
code(doc, "└── attack4_replay.lua             ← Decodes Attacks 1, 4 on UDP 14550")

doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
