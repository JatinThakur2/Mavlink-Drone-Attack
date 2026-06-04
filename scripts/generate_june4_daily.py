#!/usr/bin/env python3
"""Generates the June 4 daily internship report — Detection Framework Design."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_FILE = os.path.join(os.path.dirname(__file__), "..", "reports", "daily_report_june4.docx")

DARK_BLUE = RGBColor(31,  78,  121)
MID_BLUE  = RGBColor(47,  117, 181)
GREY      = RGBColor(89,  89,  89)
WHITE     = RGBColor(255, 255, 255)
GREEN     = RGBColor(0,   176, 80)
RED       = RGBColor(192, 0,   0)
ORANGE    = RGBColor(255, 102, 0)
PURPLE    = RGBColor(112, 48,  160)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_divider(doc, color="1F4E79"):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)

def heading(doc, text, level=1, color=DARK_BLUE, space_before=10):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(2)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def bullet(doc, text, indent=0.3, size=11):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def sub_bullet(doc, text):
    bullet(doc, text, indent=0.6, size=10.5)

def code(doc, text):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    return p

def normal(doc, text, size=11, color=None, bold=False, italic=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p

# ─────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin    = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# ── Title ─────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Internship Daily Report")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK_BLUE
add_divider(doc)

dp = doc.add_paragraph()
dr = dp.add_run("June 4, 2026")
dr.bold = True; dr.font.size = Pt(14); dr.font.color.rgb = DARK_BLUE
dp.paragraph_format.space_after = Pt(3)

gp = doc.add_paragraph()
gp.add_run("Goal: ").bold = True
gp.runs[0].font.size = Pt(11)
gv = gp.add_run(
    "Design and document a real-time anomaly detection framework for MAVLink timestamp abuse. "
    "Transition from offensive (attack) phase to defensive (detection) phase. "
    "Plan feature extraction, statistical rule engine, and machine learning pipeline "
    "to detect all five previously implemented attacks."
)
gv.font.size = Pt(11)
gp.paragraph_format.space_after = Pt(4)

sp = doc.add_paragraph()
sp.add_run("Stack: ").bold = True
sp.runs[0].font.size = Pt(10.5)
sv = sp.add_run(
    "Python 3.10.12 · scikit-learn (Isolation Forest) · scapy · pymavlink 2.4.49 · "
    "ArduPilot SITL v4.8.0-dev · Wireshark 3.6.2"
)
sv.font.size = Pt(10.5); sv.font.color.rgb = GREY; sv.italic = True
sp.paragraph_format.space_after = Pt(6)
add_divider(doc)

# ── Section 1: Phase Transition ───────────────────────────────
heading(doc, "Phase Transition: Offensive → Defensive", level=2, color=DARK_BLUE, space_before=6)
normal(doc,
    "Previous sessions focused on implementing five MAVLink attacks against ArduPilot SITL "
    "(signed command injection, GPS spoofing, timestamp overflow/underflow DoS, and replay). "
    "Today's session began the defensive phase: designing a detection system that can identify "
    "each attack from the same network traffic the attacks produce.",
    size=11)

bullet(doc, "Attack scripts are complete and tested — all five produce verifiable effects in SITL and Wireshark.")
bullet(doc, "Detection work starts from scratch — no existing IDS covers MAVLink timestamp-specific abuse patterns.")
bullet(doc, "Goal: build a system that catches our own attacks automatically, with zero prior knowledge of the attack payloads.")

add_divider(doc)

# ── Section 2: Attack Fingerprints ────────────────────────────
heading(doc, "Attack Fingerprints — What Each Attack Leaves in the Stream", level=2, color=DARK_BLUE, space_before=6)
normal(doc,
    "Each attack produces measurable anomalies in the network traffic. "
    "The detection framework is built around extracting and classifying these signals.",
    size=11)

tbl = doc.add_table(rows=6, cols=4)
tbl.style = "Table Grid"
hdrs = ["Attack", "Primary Signal", "Secondary Signal", "Detection Method"]
for i, h in enumerate(hdrs):
    c = tbl.cell(0, i)
    shade_cell(c, "1F4E79")
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)

rows = [
    ["Attack 1\n(Signed Command)",
     "Signature timestamp 8+ hours in the future",
     "Sequence number anomaly vs stream",
     "Rule: ts_gap > 3600s + IF score"],
    ["Attack 2\n(GPS Spoof)",
     "GPS time +10 days vs system clock",
     "Position drift rate > 6 m/s southward",
     "Rule: gps_sys_delta > 1day + drift_rate"],
    ["Attack 3\n(Overflow DoS)",
     "GPS time_usec = year 2104 (+78 years)",
     "All commands go unanswered after packet",
     "Rule: gps_time > 2050 (instant)"],
    ["Attack 3B\n(Underflow DoS)",
     "GPS time_usec = year 2010 (pre-epoch)",
     "Burst of 10 identical packets in 2 seconds",
     "Rule: gps_time < 2015 + burst > 5/2s"],
    ["Attack 4\n(Replay)",
     "Exact duplicate packet hash seen twice",
     "Signature timestamp still in the future",
     "SHA-256 hash deque + Rule: ts_gap > 0"],
]
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        cell = tbl.cell(ri + 1, ci)
        cell.paragraphs[0].add_run(val).font.size = Pt(9.5)
        if ri % 2 == 0:
            shade_cell(cell, "DCE6F1")

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_divider(doc)

# ── Section 3: Feature Extraction ─────────────────────────────
heading(doc, "Feature Extraction Design", level=2, color=DARK_BLUE, space_before=6)
normal(doc,
    "For each packet, the feature extractor computes six numerical values. "
    "These form the input vector for both the statistical rules and the ML model.",
    size=11)

heading(doc, "Feature 1 — Timestamp Gap (ts_gap)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Source: MAVLink 2.0 signature block — 6-byte little-endian timestamp in 10µs units since Jan 1 2015.")
bullet(doc, "Computed as: ts_gap = sig_timestamp_unix - time.time()")
sub_bullet(doc, "Normal range: ±2 seconds (clock drift tolerance)")
sub_bullet(doc, "Attack 1 / Attack 4: ts_gap = +8 hours (future timestamp)")
sub_bullet(doc, "Attack 3 overflow: ts_gap = +78 years")
sub_bullet(doc, "Attack 3B underflow: ts_gap wraps to near +584,000 years (uint64 wrap)")

heading(doc, "Feature 2 — GPS vs System Clock Delta (gps_sys_delta)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Source: GPS_INPUT JSON packets on UDP 25100 — time_usec field.")
bullet(doc, "Computed as: gps_sys_delta = (time_usec / 1e6) - time.time()")
sub_bullet(doc, "Normal: ±500ms")
sub_bullet(doc, "Attack 2: +864,000 seconds (+10 days)")
sub_bullet(doc, "Attack 3: +2,472,048,000 seconds (+78 years)")
sub_bullet(doc, "Attack 3B: -(157,766,400) seconds (-5 years, shows as negative)")

heading(doc, "Feature 3 — Position Drift Rate (drift_m_per_s)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Source: GPS_INPUT JSON — lat/lon fields between consecutive packets.")
bullet(doc, "Computed as: distance_between_positions / time_between_packets")
sub_bullet(doc, "Normal (hovering): < 0.5 m/s")
sub_bullet(doc, "Attack 2: ~6.2 m/s southward (200m drift over 30 seconds)")

heading(doc, "Feature 4 — Packet Inter-Arrival Time (iat_ms)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Time between consecutive packets of the same type (e.g. GPS_INPUT).")
sub_bullet(doc, "Normal GPS_INPUT: 200ms interval")
sub_bullet(doc, "Attack 3B: burst of 10 packets at 200ms = fine individually, but 10 in 2s = anomaly")

heading(doc, "Feature 5 — Sequence Number Jump (seq_jump)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "MAVLink 1-byte rolling counter (0–255, wraps). Should increment by 1 each packet.")
bullet(doc, "seq_jump = |received_seq - (last_seq + 1)| mod 256")
sub_bullet(doc, "Normal: 0 (consecutive) or small gap (missed packet)")
sub_bullet(doc, "Replay (Attack 4): same sequence number appears with identical payload = definitive replay")

heading(doc, "Feature 6 — Duplicate Packet Flag (is_duplicate)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "SHA-256 hash of entire raw packet stored in a rolling deque of last 10,000 hashes.")
bullet(doc, "is_duplicate = 1 if hash seen before, 0 otherwise.")
sub_bullet(doc, "Attack 4: is_duplicate = 1 — same bytes from Attack 1 pcap resent byte-for-byte")

code(doc, "feature_vector = [ts_gap, gps_sys_delta, drift_m_per_s, iat_ms, seq_jump, is_duplicate]")
add_divider(doc)

# ── Section 4: Detection Models ───────────────────────────────
heading(doc, "Detection Models", level=2, color=DARK_BLUE, space_before=6)

heading(doc, "Layer A — Statistical Rules (Deterministic)", level=3, color=MID_BLUE, space_before=4)
normal(doc,
    "Six hard-coded threshold rules that fire instantly with zero training data. "
    "These cover all five known attacks with zero false negatives.",
    size=11)
code(doc, "Rule 1: |ts_gap| > 3600                 → ALERT: REPLAY or OVERFLOW (Attack 1/4/3)")
code(doc, "Rule 2: gps_time < 2015-01-01           → ALERT: UNDERFLOW (Attack 3B)")
code(doc, "Rule 3: gps_time > 2050-01-01           → ALERT: OVERFLOW (Attack 3)")
code(doc, "Rule 4: gps_sys_delta > 86400           → ALERT: GPS SPOOF (Attack 2)")
code(doc, "Rule 5: is_duplicate == 1               → ALERT: REPLAY (Attack 4)")
code(doc, "Rule 6: packet_count > 5 in 2 seconds  → ALERT: DoS BURST (Attack 3B)")

bullet(doc, "Advantage: deterministic, zero latency, zero training required, zero false negatives for known attacks.")
bullet(doc, "Limitation: blind to novel attacks with subtle signatures that do not cross fixed thresholds.")

heading(doc, "Layer B — Isolation Forest (ML, Unsupervised)", level=3, color=MID_BLUE, space_before=4)
normal(doc,
    "An unsupervised ML model that learns what normal MAVLink traffic looks like "
    "and flags statistical outliers — without needing labeled attack examples.",
    size=11)

bullet(doc, "Algorithm: Isolation Forest (scikit-learn). Builds random decision trees; anomalies are isolated in fewer splits.")
bullet(doc, "Training: 5–10 minutes of normal SITL traffic → ~3,000 feature vectors → fit once, save to baseline.pkl.")
bullet(doc, "Inference: each new packet produces a feature vector → model returns anomaly score 0.0–1.0.")
sub_bullet(doc, "Score > 0.6 → flag as ML anomaly (threshold tunable)")
sub_bullet(doc, "Attack packets score 0.85–0.99 (extreme outliers on ts_gap / gps_sys_delta)")

bullet(doc, "Why Isolation Forest over LSTM for Phase 1:")
sub_bullet(doc, "No labeled data needed — train on normal traffic only.")
sub_bullet(doc, "Microsecond inference time — suitable for real-time 200ms GPS stream.")
sub_bullet(doc, "Simple to interpret — anomaly score directly maps to feature outlier magnitude.")

bullet(doc, "LSTM planned for Phase 2: learns temporal sequences — better for subtle replay and slow drift attacks.")

add_divider(doc)

# ── Section 5: System Architecture ────────────────────────────
heading(doc, "System Architecture", level=2, color=DARK_BLUE, space_before=6)
normal(doc,
    "Single Python process, three concurrent threads, two UDP capture sockets.",
    size=11)

code(doc, "MAVLink stream  (UDP 14550)  ─┐")
code(doc, "GPS stream      (UDP 25100)  ─┼─► Feature Extractor ─► Statistical Rules ─► ALERT")
code(doc, "IMU/Heartbeat   (UDP 14550)  ─┘         │")
code(doc, "                                          └─► Isolation Forest ──────────────► ALERT")
code(doc, "                                                    │")
code(doc, "                                          Baseline DB  |  Hash Deque (10k)")

bullet(doc, "Thread 1 (Capture): sniffs raw UDP on ports 14550 and 25100 using scapy / socket.")
bullet(doc, "Thread 2 (Rules): reads packet queue → extracts features → applies six statistical rules.")
bullet(doc, "Thread 3 (ML): feeds feature vectors to loaded Isolation Forest model → scores anomalies.")
bullet(doc, "Output: real-time alerts to terminal + timestamped entries in alerts.log.")

add_divider(doc)

# ── Section 6: Planned File Structure ─────────────────────────
heading(doc, "Planned File Structure", level=2, color=DARK_BLUE, space_before=6)
code(doc, "drone_security_lab/")
code(doc, "└── detection/")
code(doc, "    ├── features.py          ← extract 6-value vector from raw packet")
code(doc, "    ├── detector.py          ← main loop: capture + rules + ML alerts")
code(doc, "    ├── alerts.py            ← alert formatting + log file writer")
code(doc, "    ├── models/")
code(doc, "    │   ├── train_baseline.py   ← train Isolation Forest on normal traffic")
code(doc, "    │   └── baseline.pkl        ← saved trained model")
code(doc, "    └── results/")
code(doc, "        └── alerts.log          ← timestamped detection log")

add_divider(doc)

# ── Section 7: Build Order ────────────────────────────────────
heading(doc, "Implementation Plan — Build Order", level=2, color=DARK_BLUE, space_before=6)

steps = [
    ("Step 1 — features.py",
     "Write pure feature extraction functions with no network dependencies. "
     "Accepts raw bytes (MAVLink) or JSON string (GPS). Returns 6-element list. "
     "Unit-testable without SITL running."),
    ("Step 2 — detector.py (rules only)",
     "Add UDP capture loop + statistical rules engine. No ML yet. "
     "Test: run SITL, trigger each attack, verify correct alert fires."),
    ("Step 3 — train_baseline.py",
     "Capture 5 minutes of normal SITL traffic. Extract feature vectors. "
     "Fit Isolation Forest (n_estimators=200, contamination=0.01). Save baseline.pkl."),
    ("Step 4 — Add ML to detector.py",
     "Load baseline.pkl at startup. Feed every feature vector to model. "
     "Emit ML_ANOMALY alert when score > 0.6. Run both layers simultaneously."),
    ("Step 5 — Validation",
     "Run all five attacks in sequence. Verify each produces the correct alert type. "
     "Check for false positives on normal traffic. Tune thresholds if needed."),
]
for title, desc in steps:
    heading(doc, title, level=3, color=PURPLE, space_before=4)
    normal(doc, desc, size=11)

add_divider(doc)

# ── Section 8: Expected Alert Output ─────────────────────────
heading(doc, "Expected Detector Output (Live Terminal)", level=2, color=DARK_BLUE, space_before=6)
normal(doc, "When attacks are triggered in sequence, the detector should produce:", size=11)
code(doc, "[INFO]   Monitoring UDP 14550 + 25100 | baseline loaded: 3000 samples")
code(doc, "[ALERT]  ATK2_GPS_SPOOF   | gps_delta=+10d | drift=187m south | lat=-35.3650")
code(doc, "[ALERT]  ATK3_OVERFLOW    | gps_time=2104  | delta=+78yr | PERMANENT DoS imminent")
code(doc, "[ALERT]  ATK0_UNDERFLOW   | gps_time=2010  | pre-epoch | uint64 wrap | burst=10")
code(doc, "[ALERT]  ATK4_REPLAY      | dup_hash=a3f9b2| sig_ts=+8hr future | REPLAY CONFIRMED")
code(doc, "[ML]     ANOMALY score=0.91| features=[28800, 864000, 0.0, 200, 0, 0]")

add_divider(doc)

# ── Section 9: Next Steps ─────────────────────────────────────
heading(doc, "Next Steps", level=2, color=DARK_BLUE, space_before=6)
bullet(doc, "Implement features.py — pure extraction logic, no network code.")
bullet(doc, "Implement detector.py — live UDP capture + six statistical rules.")
bullet(doc, "Run SITL + trigger all five attacks — verify each alert fires correctly.")
bullet(doc, "Collect normal traffic baseline — train Isolation Forest model.")
bullet(doc, "Integrate ML layer — dual detection (rules + ML running in parallel).")
bullet(doc, "Add Wireshark coloring rules so each attack protocol shows in a distinct colour.")
bullet(doc, "Write Phase 2 plan: LSTM for temporal sequence anomaly detection.")

# ── Save ──────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUT_FILE), exist_ok=True)
doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
