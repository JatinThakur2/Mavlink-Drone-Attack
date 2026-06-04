#!/usr/bin/env python3
"""Generates the June 2 daily internship report."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_FILE = os.path.join(os.path.dirname(__file__), "..", "reports", "daily_report_june2.docx")

DARK_BLUE = RGBColor(31,  78,  121)
MID_BLUE  = RGBColor(47,  117, 181)
GREY      = RGBColor(89,  89,  89)
WHITE     = RGBColor(255, 255, 255)
GREEN     = RGBColor(0,   176, 80)
RED       = RGBColor(192, 0,   0)
ORANGE    = RGBColor(255, 102, 0)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_divider(doc, color="1F4E79"):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)

def heading(doc, text, level=1, color=DARK_BLUE, space_before=10):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(2)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def bullet(doc, text, indent=0.3, size=11):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def sub_bullet(doc, text):
    bullet(doc, text, indent=0.6, size=10.5)

def code(doc, text):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    return p

# ─────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin    = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Internship Daily Report")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK_BLUE
add_divider(doc)

# Date / Goal
dp = doc.add_paragraph()
dr = dp.add_run("June 2, 2026")
dr.bold = True; dr.font.size = Pt(14); dr.font.color.rgb = DARK_BLUE
dp.paragraph_format.space_after = Pt(3)

gp = doc.add_paragraph()
gp.add_run("Goal: ").bold = True
gp.runs[0].font.size = Pt(11)
gv = gp.add_run(
    "Align all attack scripts with the reference implementation (JSON GPS injection, "
    "manual SHA-256 signing), add Gazebo-visible position spoofing to Attack 2, "
    "implement Attack 3B (timestamp underflow), build Wireshark Lua dissectors for "
    "all attacks, and complete a full live Attack 4 replay demo in Gazebo."
)
gv.font.size = Pt(11)
gp.paragraph_format.space_after = Pt(4)

sp = doc.add_paragraph()
sp.add_run("Stack: ").bold = True
sp.runs[0].font.size = Pt(10.5)
sv = sp.add_run(
    "ArduPilot SITL v4.8.0-dev · Gazebo Harmonic · MAVProxy 1.8.74 · "
    "Wireshark 3.6.2 + Lua 5.1 · Python 3.10.12 · pymavlink 2.4.49"
)
sv.font.size = Pt(10.5); sv.font.color.rgb = GREY; sv.italic = True
sp.paragraph_format.space_after = Pt(6)
add_divider(doc)

# ── Work Completed ────────────────────────────────────────────
heading(doc, "Work Completed", level=2, color=DARK_BLUE, space_before=6)

# 1
heading(doc, "Attack Scripts Aligned with Reference Implementation", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Attack 1 — switched from pymavlink signing API to manual SHA-256 signing (matching reference custom-input.py).")
sub_bullet(doc, "Key is now hashed with SHA256 before use — same method as MAVProxy internally. Old code padded raw bytes with zeros (wrong).")
sub_bullet(doc, "Packet built with struct.pack + mavcrc.x25crc_slow + hashlib.sha256. Sent with conn.write(raw_bytes) not command_long_send().")
bullet(doc, "Attacks 2 & 3 — GPS injection switched from pymavlink binary MAVLink to JSON over raw UDP socket.")
sub_bullet(doc, "MAVProxy GPSInput module on UDP 25100 expects JSON, not binary MAVLink. Reference simulate-gps.py confirms this format.")
sub_bullet(doc, "Removed pymavlink dependency for GPS injection — just socket.sendto(json.dumps(data).encode(), (host, 25100)).")

# 2
heading(doc, "Attack 2 — Position Spoofing Added (Gazebo Visible)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Previous Attack 2 only spoofed the timestamp — drone had no reason to move, no visual effect in Gazebo.")
bullet(doc, "Added SPOOF_LAT_DRIFT = -0.0018 degrees: reported GPS position drifts 200m south over 30 seconds.")
bullet(doc, "Drone in GUIDED mode compensates for the fake southward drift by flying NORTH — visible in Gazebo.")
bullet(doc, "Two effects demonstrated simultaneously: clock +10 days (invisible) + drone moves 200m north (visible).")
sub_bullet(doc, "Verified: MAVProxy 'time' command shows 10 days in future. 'mode LAND' command rejected after attack.")

# 3
heading(doc, "Attack 3B — Timestamp Underflow DoS (New Attack)", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "New attack: inject GPS_INPUT with time_usec representing Jan 1 2010 — 5 years BEFORE the MAVLink epoch (Jan 1 2015).")
bullet(doc, "MAVLink signing timestamp = (unix - epoch) * 100000 = negative value = -15,776,640,000,000.")
bullet(doc, "Cast to uint64: wraps to 18,446,728,297,069,551,616 — larger than Attack 3's overflow value.")
bullet(doc, "Two possible effects: Effect A = DoS (same as Attack 3 via uint64 wrap) / Effect B = clock reset to 0 (accepts all future packets).")
bullet(doc, "Script: attacks/attack3b_underflow_dos.py — sends 10 burst packets, then verifies effect by sending a real signed command.")
sub_bullet(doc, "Verification: if UAV responds = clock reset (Effect B). If silent = DoS (Effect A).")

# 4
heading(doc, "Wireshark Lua Dissectors — All Attacks Decoded in GUI", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "gps_spoof.lua: decodes all GPS_INPUT packets on UDP 25100 into 3 attack zones.")
sub_bullet(doc, "Zone 0 (Attack 3B): time_usec < Jan 2015 — shows signed MAVLink TS, uint64 wrapped value, 48-bit masked value.")
sub_bullet(doc, "Zone 2 (Attack 2): time_usec in 2026 range — shows +10 day gap, lat drift in metres.")
sub_bullet(doc, "Zone 3 (Attack 3): time_usec > 3e15 — shows 78-year gap, MAVLink TS vs 2^48-1, DoS effect.")
bullet(doc, "attack4_replay.lua: post-dissector on all packets, annotates COMMAND_LONG (76) and COMMAND_ACK (77).")
sub_bullet(doc, "Shows: signature date, hours in future, command decoded as LAND, ACK result=0 (ACCEPTED).")
sub_bullet(doc, "Fixed: register_postdissector(p, true) — allfields=true makes fields indexable for display filters.")
bullet(doc, "Fixed Lua uint64 overflow: UINT64_MAX cannot be stored exactly in Lua double (max exact int = 2^53). Replaced with pre-computed string constants.")
bullet(doc, "Fixed decode_as_entries conflict: port 25100 was locked to MAVLINK_PROTO from a previous session, overriding our dissector.")

# 5
heading(doc, "Complete Attack 4 Replay Demo in Gazebo", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Full end-to-end Attack 4 sequence documented and tested in Gazebo.")
bullet(doc, "Drone flies to waypoint (-35.3619, 149.1652) at 20m AGL — visibly airborne in Gazebo.")
bullet(doc, "attack4_replay.py replays signed COMMAND_LONG from attack1_capture.pcap byte-for-byte.")
bullet(doc, "UAV accepts the replay (COMMAND_ACK result=0), switches to LAND mode, drone descends in Gazebo.")
bullet(doc, "Key condition: SITL rebooted (clock reset) but key NOT rotated — old future timestamp still valid.")
bullet(doc, "Timestamp validity check script added — warns if attack1 capture timestamp has expired (>24h old).")

# 6
heading(doc, "Live Wireshark Capture Setup", level=3, color=MID_BLUE, space_before=4)
bullet(doc, "Documented live capture on loopback (lo) interface with capture filter: udp port 25100 or udp port 14550.")
bullet(doc, "Dissectors apply automatically to live traffic — same decoded fields appear in real-time during attacks.")
bullet(doc, "Confirmed: decode_as_entries persists across sessions — port 25100 → atk2gps, port 14550 → mavlink_proto.")

add_divider(doc)

# ── Issues & Resolutions ──────────────────────────────────────
heading(doc, "Issues Faced & Resolutions", level=2, color=DARK_BLUE, space_before=6)

bullet(doc, "Signing key derivation mismatch — old code used raw bytes padded to 32 bytes; MAVProxy uses sha256(key_string).")
sub_bullet(doc, "Resolution: replaced ljust(32, b'\\x00') with hashlib.sha256(key_str.encode()).digest().")

bullet(doc, "GPS_INPUT format — old code sent binary MAVLink packets; GPSInput module expects JSON.")
sub_bullet(doc, "Resolution: replaced mavutil.mavlink_connection + gps_input_send() with raw UDP socket + json.dumps().")

bullet(doc, "No visual effect in Gazebo for Attack 2 — timestamp spoofing alone doesn't move the drone.")
sub_bullet(doc, "Resolution: added gradual lat drift (-0.0018 degrees over 30s) — drone compensates northward, fully visible.")

bullet(doc, "Lua uint64 overflow — UINT64_MAX = 2^64-1 cannot be stored exactly in Lua double precision (max 2^53).")
sub_bullet(doc, "Resolution: pre-computed Attack 3B wrapped values stored as string constants. Only 48-bit masked value (< 2^53) kept as number.")

bullet(doc, "Wireshark filter atk4replay.type returned 0 results — post-dissector fields not indexed by default.")
sub_bullet(doc, "Resolution: changed register_postdissector(p) to register_postdissector(p, true) to enable field indexing.")

bullet(doc, "Port 25100 locked to MAVLINK_PROTO in decode_as_entries from a previous session.")
sub_bullet(doc, "Resolution: edited ~/.config/wireshark/decode_as_entries to replace MAVLINK_PROTO with atk2gps.")

bullet(doc, "Drone auto-disarming during flight — arm throttle rejected pre-arm checks (GPS/EKF).")
sub_bullet(doc, "Resolution: use arm throttle force to bypass pre-arm checks for SITL testing.")

bullet(doc, "SIM_GPS1_DISABLE parameter not found — not available in JSON SITL model.")
sub_bullet(doc, "Resolution: parameter not needed with JSON model — GPS1_TYPE=14 is sufficient for GPSInput to take priority.")
add_divider(doc)

# ── Results Table ─────────────────────────────────────────────
heading(doc, "Attack Results Summary", level=2, color=DARK_BLUE, space_before=6)

tbl = doc.add_table(rows=6, cols=5)
tbl.style = "Table Grid"

headers = ["Attack", "Method", "Visible in Gazebo", "Wireshark Protocol", "Pcap"]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    shade_cell(cell, "1F4E79")
    run = cell.paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(9.5)

rows_data = [
    ("1 — Clock Manip.",    "Manual SHA-256 signed CMD",   "Drone lands mid-flight",          "ATK4_REPLAY",      "attack1_capture.pcap"),
    ("2 — GPS Spoof",       "JSON GPS +10d + lat drift",   "Drone flies 200m north",          "ATK2_GPS_SPOOF",   "attack2_capture.pcap"),
    ("3 — Overflow DoS",    "JSON GPS at year 2104",       "Drone ignores all commands",      "ATK3_OVERFLOW",    "attack3_capture.pcap"),
    ("3B — Underflow DoS",  "JSON GPS at year 2010",       "DoS or clock reset (impl. dep.)", "ATK0_UNDERFLOW",   "attack3b_capture.pcap"),
    ("4 — Replay",          "Raw bytes from Attack 1",     "Drone lands again after reboot",  "ATK4_ACCEPTED",    "attack4_capture.pcap"),
]
for ri, row_data in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    bg  = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row_data):
        shade_cell(row.cells[ci], bg)
        row.cells[ci].paragraphs[0].add_run(val).font.size = Pt(9)

doc.add_paragraph()
add_divider(doc)

# ── Wireshark Filter Reference ────────────────────────────────
heading(doc, "Wireshark Filter Reference", level=2, color=DARK_BLUE, space_before=6)

tbl2 = doc.add_table(rows=6, cols=3)
tbl2.style = "Table Grid"

h2 = ["Filter", "Shows", "Attack"]
for i, h in enumerate(h2):
    cell = tbl2.rows[0].cells[i]
    shade_cell(cell, "1F4E79")
    run = cell.paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(9.5)

filters = [
    ("udp.dstport == 25100",                           "All GPS injection packets",            "2, 3, 3B"),
    ("mavlink_proto.msgid == 76 || mavlink_proto.msgid == 77", "COMMAND_LONG + COMMAND_ACK",  "1, 4"),
    ("atk4replay.type",                                "Only replay & ACK packets",            "1, 4"),
    ("frame.protocols contains \"atk4replay\"",        "All packets our dissector touched",    "1, 4"),
    ("atk2gps.gap",                                    "Packets with 10-day clock gap",        "2"),
]
for ri, row_data in enumerate(filters):
    row = tbl2.rows[ri + 1]
    bg  = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row_data):
        shade_cell(row.cells[ci], bg)
        run = row.cells[ci].paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if ci == 0:
            run.font.name = "Courier New"

doc.add_paragraph()
add_divider(doc)

# ── Next Steps ────────────────────────────────────────────────
heading(doc, "Next Steps", level=2, color=DARK_BLUE, space_before=6)
bullet(doc, "Run all 5 attacks in a single session with live Wireshark open — capture each pcap fresh.")
bullet(doc, "Add Wireshark coloring rules so each attack protocol shows in a distinct colour.")
bullet(doc, "Update the full technical report (generate_full_report_v2.py) to include Attack 3B and all dissector details.")
bullet(doc, "Document the arm throttle force workaround and SIM_GPS1_DISABLE findings in the project README.")

doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
